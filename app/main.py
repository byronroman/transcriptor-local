from __future__ import annotations

import hashlib
import json
import math
import os
import platform
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import unicodedata
import uuid
import wave
import webbrowser
import zipfile
from collections import Counter
from contextlib import asynccontextmanager
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Any, Callable, Optional

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, PlainTextResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel


ROOT_DIR = Path(__file__).resolve().parents[1]
STATIC_DIR = ROOT_DIR / "static"
DATA_DIR = ROOT_DIR / "data"
PROJECTS_DIR = DATA_DIR / "projects"
EXPORTS_DIR = DATA_DIR / "exports"
MODELS_DIR = ROOT_DIR / "models"
TOOLS_DIR = ROOT_DIR / "tools"
LANGUAGETOOL_DIR = TOOLS_DIR / "languagetool"
LANGUAGETOOL_VERSION = os.environ.get("LANGUAGETOOL_VERSION", "6.6")
MIN_LANGUAGETOOL_JAVA_MAJOR = int(os.environ.get("MIN_LANGUAGETOOL_JAVA_MAJOR", "17"))
LANGUAGETOOL_JAR_PATTERNS = (
    "languagetool-server.jar",
    "languagetool-standalone*.jar",
    "LanguageTool.jar",
    "LanguageTool.uno.jar",
)
MAX_IMPORT_PACKAGE_BYTES = int(os.environ.get("MAX_IMPORT_PACKAGE_BYTES", str(2_200_000_000)))
MAX_PACKAGE_MEMBERS = int(os.environ.get("MAX_PACKAGE_MEMBERS", "32"))
MAX_PACKAGE_JSON_BYTES = int(os.environ.get("MAX_PACKAGE_JSON_BYTES", str(50_000_000)))
MAX_PACKAGE_DIAGNOSTIC_BYTES = int(os.environ.get("MAX_PACKAGE_DIAGNOSTIC_BYTES", str(50_000_000)))
MAX_PACKAGE_AUDIO_BYTES = int(os.environ.get("MAX_PACKAGE_AUDIO_BYTES", str(2_000_000_000)))
MAX_UPLOAD_AUDIO_BYTES = int(os.environ.get("MAX_UPLOAD_AUDIO_BYTES", str(2_200_000_000)))
MAX_IMPORTED_SEGMENTS = int(os.environ.get("MAX_IMPORTED_SEGMENTS", "100000"))
MAX_IMPORTED_TEXT_CHARS = int(os.environ.get("MAX_IMPORTED_TEXT_CHARS", "50000"))
EXPORT_TEMP_MAX_AGE_SECONDS = int(os.environ.get("EXPORT_TEMP_MAX_AGE_SECONDS", str(24 * 60 * 60)))
EXPORT_TEMP_PREFIX = ".tmp-"
ALLOWED_PACKAGE_DIAGNOSTICS = {
    "whisper_quality.json",
    "diarization_quality.json",
    "diarization_turns.json",
    "process.log",
}
ALLOWED_IMPORT_AUDIO_SUFFIXES = {
    ".aac",
    ".aiff",
    ".aif",
    ".audio",
    ".flac",
    ".m4a",
    ".mkv",
    ".mov",
    ".mp3",
    ".mp4",
    ".ogg",
    ".opus",
    ".wav",
    ".webm",
    ".wma",
}

HOST = "127.0.0.1"
PORT = int(os.environ.get("TRANSCRIPTOR_PORT", "8765"))
DIARIZATION_CHUNK_SECONDS = float(os.environ.get("DIARIZATION_CHUNK_SECONDS", "180"))
DIARIZATION_CHUNK_OVERLAP_SECONDS = float(os.environ.get("DIARIZATION_CHUNK_OVERLAP_SECONDS", "3"))
DEFAULT_DIARIZATION_SPEAKERS = int(os.environ.get("DEFAULT_DIARIZATION_SPEAKERS", "2"))
WHISPER_CHUNK_SECONDS = float(os.environ.get("WHISPER_CHUNK_SECONDS", "120"))
WHISPER_CHUNK_OVERLAP_SECONDS = float(os.environ.get("WHISPER_CHUNK_OVERLAP_SECONDS", "10"))
WHISPER_RETRY_CHUNK_SECONDS = tuple(
    float(value.strip())
    for value in os.environ.get("WHISPER_RETRY_CHUNK_SECONDS", "60,45,30,20").split(",")
    if value.strip()
)
WHISPER_MAX_NON_SILENT_GAP_SECONDS = float(os.environ.get("WHISPER_MAX_NON_SILENT_GAP_SECONDS", "30"))
WHISPER_RMS_SPEECH_THRESHOLD = float(os.environ.get("WHISPER_RMS_SPEECH_THRESHOLD", "80"))
WHISPER_MAX_LEN = int(os.environ.get("WHISPER_MAX_LEN", "120"))

PROCESSING_PROFILES: dict[str, dict[str, Any]] = {
    "calidad": {
        "label": "Calidad",
        "chunk_seconds": 120.0,
        "overlap_seconds": 10.0,
        "retry_chunk_seconds": [60.0, 45.0, 30.0, 20.0],
        "max_threads": 4,
    },
    "seguro_windows": {
        "label": "Seguro Windows",
        "chunk_seconds": 60.0,
        "overlap_seconds": 8.0,
        "retry_chunk_seconds": [45.0, 30.0, 20.0],
        "max_threads": 1,
    },
    "rapido": {
        "label": "Rapido",
        "chunk_seconds": 90.0,
        "overlap_seconds": 8.0,
        "retry_chunk_seconds": [45.0, 30.0, 20.0],
        "max_threads": 4,
        "prefer_model": "large-v3-turbo",
    },
}

WHISPER_OPTION_CACHE: dict[str, set[str]] = {}
MODEL_INTEGRITY_CACHE: dict[tuple[str, int, int], Optional[str]] = {}
WHISPER_MODEL_INTEGRITY = {
    "ggml-large-v3.bin": {
        "size": 3095033483,
        "sha256": "64d182b440b98d5203c4f9bd541544d84c605196c4f7b845dfa11fb23594d1e2",
    }
}
PROJECT_ID_RE = re.compile(r"^[a-f0-9]{12}$")
PROJECT_NOT_FOUND_DETAIL = "Proyecto no encontrado"

for directory in (DATA_DIR, PROJECTS_DIR, EXPORTS_DIR, MODELS_DIR, TOOLS_DIR, LANGUAGETOOL_DIR):
    directory.mkdir(parents=True, exist_ok=True)


@asynccontextmanager
async def lifespan(_: FastAPI):
    normalize_stale_projects()
    yield


app = FastAPI(title="Transcriptor Mi Cami", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[f"http://{HOST}:{PORT}", "http://localhost:8765"],
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

JOBS: dict[str, dict[str, Any]] = {}
RUNNING_PROCESSES: dict[str, subprocess.Popen[str]] = {}
JOBS_LOCK = threading.Lock()
PROJECT_FILE_LOCKS: dict[str, threading.Lock] = {}
PROJECT_FILE_LOCKS_LOCK = threading.Lock()
PROOFREAD_TOOL: Any = None
PROOFREAD_LOCK = threading.Lock()
PROOFREAD_STATUS_LOCK = threading.Lock()
PROOFREAD_CACHE_LOCK = threading.Lock()
PROOFREAD_INIT_STARTED = False
PROOFREAD_STATUS: dict[str, Any] = {
    "status": "idle",
    "available": False,
    "local": True,
    "language": "es",
    "message": "Corrector local sin iniciar.",
    "missing": [],
    "updated_at": None,
}
PROOFREAD_CACHE: dict[str, dict[str, Any]] = {}
PROOFREAD_CACHE_ORDER: list[str] = []
PROOFREAD_CACHE_LIMIT = 1200


class SegmentUpdate(BaseModel):
    segments: list[dict[str, Any]]
    speaker_labels: dict[str, str] = {}
    name: Optional[str] = None
    base_content_revision: Optional[int] = None


class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    playback_position: Optional[float] = None


class RelabelRequest(BaseModel):
    mode: str = "interview_2p"


class ProofreadRequest(BaseModel):
    text: str
    language: str = "es"


class ProofreadBatchItem(BaseModel):
    id: str
    text: str


class ProofreadBatchRequest(BaseModel):
    items: list[ProofreadBatchItem]
    language: str = "es"


@dataclass
class ToolPaths:
    ffmpeg: Optional[str]
    whisper: Optional[str]


class JobStopped(RuntimeError):
    def __init__(self, status: str):
        self.status = status
        message = "Proceso pausado" if status == "paused" else "Proceso cancelado"
        super().__init__(message)


def now_ms() -> int:
    return int(time.time() * 1000)


def clean_filename(name: str) -> str:
    stem = Path(name).stem or "audio"
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "", stem).strip()
    return stem[:80] or "audio"


def validate_project_id(project_id: str) -> str:
    value = str(project_id or "")
    if value != value.strip():
        raise HTTPException(status_code=404, detail=PROJECT_NOT_FOUND_DETAIL)
    value = value.lower()
    if not PROJECT_ID_RE.fullmatch(value):
        raise HTTPException(status_code=404, detail=PROJECT_NOT_FOUND_DETAIL)
    return value


def safe_project_dir(project_id: str) -> Path:
    validated = validate_project_id(project_id)
    root = PROJECTS_DIR.resolve()
    path = (root / validated).resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=PROJECT_NOT_FOUND_DETAIL) from exc
    return path


def project_dir(project_id: str) -> Path:
    return safe_project_dir(project_id)


def project_owned_path(project_id: str, raw_path: Any) -> Optional[Path]:
    text = str(raw_path or "").strip()
    if not text:
        return None
    pdir = project_dir(project_id).resolve()
    candidate = Path(text)
    if not candidate.is_absolute():
        candidate = pdir / candidate
    try:
        resolved = candidate.resolve()
        resolved.relative_to(pdir)
    except (OSError, ValueError):
        return None
    return resolved


def path_leaf(raw_path: Any) -> str:
    text = str(raw_path or "").strip()
    if not text or "\x00" in text:
        return ""
    return PurePosixPath(text.replace("\\", "/")).name


def allowed_audio_file(path: Path) -> bool:
    return path.is_file() and (path.suffix.lower() or ".audio") in ALLOWED_IMPORT_AUDIO_SUFFIXES


def project_media_fallback(project_id: str, raw_path: Any, role: str, source_name: Any = "") -> Optional[Path]:
    pdir = project_dir(project_id)
    candidates: list[Path] = []
    seen: set[str] = set()

    def add_name(name: Any) -> None:
        leaf = path_leaf(name)
        if not leaf:
            return
        candidate = pdir / leaf
        key = candidate.name.lower()
        if key not in seen:
            seen.add(key)
            candidates.append(candidate)

    def add_path(path: Path) -> None:
        key = path.name.lower()
        if key not in seen:
            seen.add(key)
            candidates.append(path)

    add_name(raw_path)
    if role == "source":
        add_name(source_name)
        for path in sorted(pdir.glob("source.*")):
            add_path(path)
    elif role == "audio":
        add_name("audio_16k.wav")
        for path in sorted(pdir.glob("audio_16k.*")):
            add_path(path)

    for candidate in candidates:
        if allowed_audio_file(candidate):
            return candidate
    return None


def project_media_path(project: dict[str, Any], key: str, role: str) -> Optional[Path]:
    project_id = validate_project_id(str(project.get("id") or ""))
    owned = project_owned_path(project_id, project.get(key))
    if owned and allowed_audio_file(owned):
        return owned
    return project_media_fallback(project_id, project.get(key), role, project.get("source_name"))


def repair_project_media_paths(project: dict[str, Any]) -> list[str]:
    project_id = validate_project_id(str(project.get("id") or ""))
    pdir = project_dir(project_id).resolve()
    repaired: list[str] = []
    for key, role in (("source_path", "source"), ("audio_path", "audio")):
        value = project.get(key)
        if not value:
            continue
        owned = project_owned_path(project_id, value)
        if owned and allowed_audio_file(owned):
            continue
        fallback = project_media_fallback(project_id, value, role, project.get("source_name"))
        if not fallback:
            continue
        try:
            project[key] = str(fallback.resolve().relative_to(pdir))
        except ValueError:
            continue
        repaired.append(key)
    return repaired


def metadata_path(project_id: str) -> Path:
    return project_dir(project_id) / "project.json"


def project_file_lock(project_id: str) -> threading.Lock:
    project_id = validate_project_id(project_id)
    with PROJECT_FILE_LOCKS_LOCK:
        if project_id not in PROJECT_FILE_LOCKS:
            PROJECT_FILE_LOCKS[project_id] = threading.Lock()
        return PROJECT_FILE_LOCKS[project_id]


def recover_project_json(path: Path, text: str, error: json.JSONDecodeError) -> Optional[dict[str, Any]]:
    if error.msg != "Extra data":
        return None
    decoder = json.JSONDecoder()
    try:
        project, end = decoder.raw_decode(text)
    except json.JSONDecodeError:
        return None
    extra = text[end:]
    if extra.strip().strip("}") != "":
        return None
    backup = path.with_name(f"{path.name}.corrupt-{int(time.time())}")
    try:
        shutil.copy2(path, backup)
    except OSError:
        pass
    project["id"] = validate_project_id(path.parent.name)
    save_project(project)
    return project


def load_project(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    path = metadata_path(project_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    text = path.read_text(encoding="utf-8")
    try:
        project = json.loads(text)
    except json.JSONDecodeError as error:
        recovered = recover_project_json(path, text, error)
        if recovered is not None:
            append_project_log(project_id, "project.json tenia datos extra al final; se creo backup y se reparo automaticamente.")
            project = recovered
        else:
            raise HTTPException(status_code=500, detail=f"project.json corrupto: {error}") from error
    project["id"] = project_id
    repaired_paths = repair_project_media_paths(project)
    if repaired_paths:
        project["updated_at"] = now_ms()
        save_project(project)
        append_project_log(project_id, f"Rutas locales reparadas tras migracion: {', '.join(repaired_paths)}.")
    project["content_revision"] = project_content_revision(project)
    return project


def save_project(project: dict[str, Any]) -> None:
    project["id"] = validate_project_id(str(project["id"]))
    path = metadata_path(project["id"])
    lock = project_file_lock(str(project["id"]))
    with lock:
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_name(f"{path.name}.{os.getpid()}.{threading.get_ident()}.tmp")
        try:
            tmp.write_text(json.dumps(project, ensure_ascii=False, indent=2), encoding="utf-8")
            tmp.replace(path)
        finally:
            tmp.unlink(missing_ok=True)


def project_content_revision(project: dict[str, Any]) -> int:
    try:
        value = int(project.get("content_revision", 0) or 0)
    except (TypeError, ValueError):
        return 0
    return max(0, value)


def bump_content_revision(project: dict[str, Any]) -> int:
    next_revision = project_content_revision(project) + 1
    project["content_revision"] = next_revision
    return next_revision


def log_path(project_id: str) -> Path:
    return project_dir(project_id) / "process.log"


def processing_manifest_path(project_id: str) -> Path:
    return project_dir(project_id) / "processing_manifest.json"


def diarization_manifest_path(project_id: str) -> Path:
    return project_dir(project_id) / "diarization_manifest.json"


def write_json_atomic(path: Path, payload: dict[str, Any] | list[Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(f"{path.name}.{os.getpid()}.{threading.get_ident()}.tmp")
    try:
        tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)
    finally:
        tmp.unlink(missing_ok=True)


def exports_root() -> Path:
    root = EXPORTS_DIR.resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def safe_export_path(project_id: str, suffix: str) -> Path:
    project_id = validate_project_id(project_id)
    suffix = str(suffix or "")
    if not suffix or "\x00" in suffix or "/" in suffix or "\\" in suffix:
        raise HTTPException(status_code=400, detail="Nombre de exportacion invalido.")
    root = exports_root()
    path = (root / f"{project_id}{suffix}").resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Ruta de exportacion invalida.") from exc
    return path


def temporary_export_path(label: str, suffix: str = ".tmp") -> Path:
    label = re.sub(r"[^A-Za-z0-9_-]+", "-", str(label or "export")).strip("-") or "export"
    suffix = str(suffix or ".tmp")
    if "\x00" in suffix or "/" in suffix or "\\" in suffix:
        raise HTTPException(status_code=400, detail="Nombre temporal invalido.")
    return exports_root() / f"{EXPORT_TEMP_PREFIX}{label}-{uuid.uuid4().hex}{suffix}"


def cleanup_stale_export_temps(max_age_seconds: int = EXPORT_TEMP_MAX_AGE_SECONDS) -> None:
    cutoff = time.time() - max(60, int(max_age_seconds))
    for path in exports_root().glob(f"{EXPORT_TEMP_PREFIX}*"):
        try:
            if path.is_file() and path.stat().st_mtime < cutoff:
                path.unlink(missing_ok=True)
        except OSError:
            continue


def write_export_atomically(output_path: Path, writer: Callable[[Path], None]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    cleanup_stale_export_temps()
    tmp = temporary_export_path(output_path.stem, output_path.suffix or ".tmp")
    try:
        writer(tmp)
        tmp.replace(output_path)
    except Exception:
        tmp.unlink(missing_ok=True)
        raise


def read_json_or_none(path: Path) -> Optional[Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (FileNotFoundError, json.JSONDecodeError):
        return None


def update_processing_manifest(project_id: str, **values: Any) -> None:
    path = processing_manifest_path(project_id)
    payload = read_json_or_none(path)
    if not isinstance(payload, dict):
        return
    payload.update(values)
    payload["updated_at"] = now_ms()
    write_json_atomic(path, payload)


def append_project_log(project_id: str, message: str) -> None:
    line = f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] {message.rstrip()}\n"
    path = log_path(project_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as out:
        out.write(line)


def read_project_log(project_id: str, limit: int = 12000) -> str:
    path = log_path(project_id)
    if not path.exists():
        return ""
    text = path.read_text(encoding="utf-8", errors="replace")
    return text[-limit:]


def project_summary(project: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": project.get("id"),
        "name": project.get("name"),
        "status": project.get("status"),
        "created_at": project.get("created_at"),
        "updated_at": project.get("updated_at"),
        "segments": len(project.get("segments") or []),
        "speakers": len(project.get("speaker_labels") or {}),
    }


def list_projects() -> list[dict[str, Any]]:
    normalize_stale_projects()
    projects = []
    for path in sorted(PROJECTS_DIR.glob("*/project.json"), reverse=True):
        try:
            project_id = validate_project_id(path.parent.name)
            project = json.loads(path.read_text(encoding="utf-8"))
            project["id"] = project_id
            projects.append(project)
        except json.JSONDecodeError:
            continue
        except HTTPException:
            continue
    return sorted(projects, key=lambda item: item.get("created_at", 0), reverse=True)


def normalize_stale_projects() -> None:
    stale_statuses = {"queued", "processing", "pausing", "cancelling"}
    for path in PROJECTS_DIR.glob("*/project.json"):
        try:
            project = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            continue
        try:
            project_id = validate_project_id(path.parent.name)
        except HTTPException:
            continue
        project["id"] = project_id
        with JOBS_LOCK:
            has_job = project_id in JOBS or project_id in RUNNING_PROCESSES
        if project.get("status") not in stale_statuses or has_job:
            continue
        warnings = project.get("warnings") or []
        message = "Proceso anterior interrumpido al cerrar o recargar la app. Puedes reanudar o eliminar el proyecto."
        if message not in warnings:
            warnings.append(message)
        project["status"] = "paused"
        project["warnings"] = warnings
        project["error"] = None
        project["updated_at"] = now_ms()
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(project, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)


def which_local(candidates: list[Path], system_names: list[str]) -> Optional[str]:
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    for name in system_names:
        found = shutil.which(name)
        if found:
            return found
    return None


def get_tools() -> ToolPaths:
    ffmpeg_candidates = [
        TOOLS_DIR / "ffmpeg" / "ffmpeg.exe",
        TOOLS_DIR / "ffmpeg" / "ffmpeg",
        TOOLS_DIR / "ffmpeg.exe",
    ]
    whisper_candidates = [
        TOOLS_DIR / "whisper" / "whisper-cli.exe",
        TOOLS_DIR / "whisper" / "whisper-cli",
        TOOLS_DIR / "whisper" / "main.exe",
        TOOLS_DIR / "whisper" / "main",
    ]
    return ToolPaths(
        ffmpeg=which_local(ffmpeg_candidates, ["ffmpeg"]),
        whisper=which_local(whisper_candidates, ["whisper-cli", "whisper-cpp", "main"]),
    )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def default_processing_profile() -> str:
    return "calidad" if platform.system() == "Darwin" else "seguro_windows"


def normalize_processing_profile(profile: Optional[str]) -> str:
    value = (profile or "auto").strip().lower()
    if value == "auto":
        return default_processing_profile()
    if value not in PROCESSING_PROFILES:
        return default_processing_profile()
    return value


def processing_profile_settings(profile: Optional[str]) -> dict[str, Any]:
    normalized = normalize_processing_profile(profile)
    settings = dict(PROCESSING_PROFILES[normalized])
    settings["name"] = normalized
    settings["retry_chunk_seconds"] = list(settings.get("retry_chunk_seconds") or [])
    return settings


def select_model_for_request(models: list[dict[str, Any]], requested_model: str, profile: str) -> str:
    if not models:
        return "auto"
    if requested_model != "auto" and any(item["name"] == requested_model for item in models):
        return requested_model
    prefer = PROCESSING_PROFILES.get(profile, {}).get("prefer_model")
    if prefer:
        for item in models:
            if str(prefer) in str(item.get("name", "")):
                return str(item["name"])
    return str(models[0]["name"])


def audio_signature(path: Path, duration: float) -> dict[str, Any]:
    stat = path.stat()
    return {
        "path": str(path),
        "size": stat.st_size,
        "sha256": sha256_file(path),
        "duration": round(duration, 3),
    }


def model_integrity_error(path: Path) -> Optional[str]:
    expected = WHISPER_MODEL_INTEGRITY.get(path.name)
    if not expected or not path.exists():
        return None
    stat = path.stat()
    expected_size = int(expected.get("size") or 0)
    if expected_size and stat.st_size != expected_size:
        return f"tamano esperado {expected_size} bytes, encontrado {stat.st_size} bytes"
    cache_key = (str(path), stat.st_mtime_ns, stat.st_size)
    if cache_key in MODEL_INTEGRITY_CACHE:
        return MODEL_INTEGRITY_CACHE[cache_key]
    expected_hash = str(expected.get("sha256") or "")
    error = None
    if expected_hash:
        actual_hash = sha256_file(path)
        if actual_hash != expected_hash:
            error = f"sha256 esperado {expected_hash}, encontrado {actual_hash}"
    MODEL_INTEGRITY_CACHE[cache_key] = error
    return error


def available_models() -> list[dict[str, Any]]:
    model_dir = MODELS_DIR / "whisper"
    model_dir.mkdir(parents=True, exist_ok=True)
    paths = sorted(
        (path for path in model_dir.glob("*.bin") if not model_integrity_error(path)),
        key=model_rank,
    )
    models = []
    for index, path in enumerate(paths):
        size_mb = path.stat().st_size / (1024 * 1024)
        models.append(
            {
                "name": path.name,
                "path": str(path),
                "size_mb": round(size_mb, 1),
                "recommended": index == 0,
            }
        )
    return models


def model_rank(path: Path) -> tuple[int, str]:
    name = path.name
    is_mac_apple = platform.system() == "Darwin" and platform.machine() == "arm64"
    if is_mac_apple:
        order = [
            "ggml-large-v3.bin",
            "large-v3-q5_0",
            "large-v3-turbo",
            "medium",
            "small",
        ]
    else:
        order = [
            "ggml-large-v3.bin",
            "large-v3-q5_0",
            "large-v3-turbo",
            "medium",
            "small",
        ]
    for index, marker in enumerate(order):
        if marker in name:
            return index, name
    return len(order), name


def diarization_ready() -> dict[str, Any]:
    segmentation = MODELS_DIR / "diarization" / "sherpa-onnx-pyannote-segmentation-3-0" / "model.onnx"
    embedding = MODELS_DIR / "diarization" / "3dspeaker_speech_eres2net_base_sv_zh-cn_3dspeaker_16k.onnx"
    package_ok = True
    error = None
    try:
        import sherpa_onnx  # noqa: F401
    except Exception as exc:  # pragma: no cover - environment dependent
        package_ok = False
        error = str(exc)
    return {
        "package": package_ok,
        "segmentation_model": segmentation.exists(),
        "embedding_model": embedding.exists(),
        "ready": package_ok and segmentation.exists() and embedding.exists(),
        "error": error,
    }


def update_job(project_id: str, **values: Any) -> None:
    with JOBS_LOCK:
        current = JOBS.setdefault(project_id, {})
        current.update(values)
        status = values.get("status", current.get("status", ""))
    if "step" in values:
        append_project_log(project_id, f"JOB {status}: {values['step']}")


def public_job(job: dict[str, Any]) -> dict[str, Any]:
    return {
        key: value
        for key, value in job.items()
        if key not in {"process"} and isinstance(value, (str, int, float, bool, type(None)))
    }


def get_stop_request(project_id: str) -> Optional[str]:
    with JOBS_LOCK:
        value = JOBS.get(project_id, {}).get("stop_requested")
        return str(value) if value else None


def clear_stop_request(project_id: str) -> None:
    with JOBS_LOCK:
        if project_id in JOBS:
            JOBS[project_id].pop("stop_requested", None)


def ensure_not_stopped(project_id: str) -> None:
    requested = get_stop_request(project_id)
    if requested in {"paused", "cancelled"}:
        raise JobStopped(requested)


def request_job_stop(project_id: str, status: str) -> dict[str, Any]:
    if status not in {"paused", "cancelled"}:
        raise HTTPException(status_code=400, detail="Estado de detencion invalido")
    project = load_project(project_id)
    if project.get("status") not in {"queued", "processing"}:
        return {"ok": True, "status": project.get("status")}

    ui_status = "pausing" if status == "paused" else "cancelling"
    update_job(
        project_id,
        status=ui_status,
        step="Pausando..." if status == "paused" else "Cancelando...",
        stop_requested=status,
    )
    with JOBS_LOCK:
        process = RUNNING_PROCESSES.get(project_id)
    if process and process.poll() is None:
        append_project_log(project_id, f"Terminando proceso activo con estado solicitado: {status}")
        process.terminate()
    return {"ok": True, "status": ui_status}


def stop_project_for_delete(project_id: str) -> None:
    with JOBS_LOCK:
        process = RUNNING_PROCESSES.get(project_id)
        if project_id in JOBS:
            JOBS[project_id]["stop_requested"] = "cancelled"
            JOBS[project_id]["status"] = "cancelling"
            JOBS[project_id]["step"] = "Eliminando proyecto..."
    if process and process.poll() is None:
        process.terminate()
        try:
            process.wait(timeout=5)
        except subprocess.TimeoutExpired:
            process.kill()
            process.wait(timeout=5)


def run_command(project_id: str, command: list[str], cwd: Optional[Path] = None) -> subprocess.CompletedProcess[str]:
    ensure_not_stopped(project_id)
    append_project_log(project_id, "RUN " + " ".join(command))
    process = subprocess.Popen(
        command,
        cwd=str(cwd) if cwd else None,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
    )
    with JOBS_LOCK:
        RUNNING_PROCESSES[project_id] = process
    try:
        output, _ = process.communicate()
    finally:
        with JOBS_LOCK:
            if RUNNING_PROCESSES.get(project_id) is process:
                RUNNING_PROCESSES.pop(project_id, None)
    requested = get_stop_request(project_id)
    if requested in {"paused", "cancelled"}:
        append_project_log(project_id, f"STOP {requested}")
        raise JobStopped(requested)
    append_project_log(project_id, f"EXIT {process.returncode}")
    if output:
        append_project_log(project_id, output[-4000:])
    return subprocess.CompletedProcess(command, process.returncode, output or "", None)


def parse_time_value(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value) / 1000.0 if value > 10000 else float(value)
    if not isinstance(value, str):
        return 0.0
    text = value.strip().replace(",", ".")
    parts = text.split(":")
    try:
        if len(parts) == 3:
            return int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2])
        if len(parts) == 2:
            return int(parts[0]) * 60 + float(parts[1])
        return float(text)
    except ValueError:
        return 0.0


def parse_whisper_time(offset_value: Any, timestamp_value: Any, fallback: Any = 0) -> float:
    if isinstance(offset_value, (int, float)):
        return float(offset_value) / 1000.0
    if isinstance(offset_value, str) and offset_value.strip().isdigit():
        return float(offset_value.strip()) / 1000.0
    return parse_time_value(timestamp_value if timestamp_value is not None else fallback)


def parse_whisper_json(path: Path) -> list[dict[str, Any]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    raw_segments = data.get("transcription") or data.get("segments") or []
    segments = []
    for idx, item in enumerate(raw_segments):
        timestamps = item.get("timestamps") or {}
        offsets = item.get("offsets") or {}
        start = parse_whisper_time(offsets.get("from"), timestamps.get("from"), item.get("start", 0))
        end = parse_whisper_time(offsets.get("to"), timestamps.get("to"), item.get("end", start))
        text = (item.get("text") or "").strip()
        if not text:
            continue
        segments.append(
            {
                "id": f"seg-{idx:05d}",
                "start": round(start, 3),
                "end": round(max(end, start), 3),
                "speaker": "SPEAKER_00",
                "text": text,
            }
        )
    return segments


TIMESTAMP_RE = re.compile(
    r"\[(?P<start>\d{2}:\d{2}:\d{2}[,.]\d{3})\s*-->\s*(?P<end>\d{2}:\d{2}:\d{2}[,.]\d{3})\]\s*(?P<text>.*)"
)
DIARIZE_PROGRESS_RE = re.compile(r"^DIARIZE_PROGRESS\s+(?P<percent>\d+(?:\.\d+)?)\s*(?P<message>.*)$")


def parse_whisper_stdout(text: str) -> list[dict[str, Any]]:
    segments = []
    for line in text.splitlines():
        match = TIMESTAMP_RE.search(line)
        if not match:
            continue
        body = match.group("text").strip()
        if not body:
            continue
        idx = len(segments)
        segments.append(
            {
                "id": f"seg-{idx:05d}",
                "start": round(parse_time_value(match.group("start")), 3),
                "end": round(parse_time_value(match.group("end")), 3),
                "speaker": "SPEAKER_00",
                "text": body,
            }
        )
    return segments


def parse_speaker_count(value: str) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return DEFAULT_DIARIZATION_SPEAKERS
    return parsed if parsed > 0 else DEFAULT_DIARIZATION_SPEAKERS


def normalize_repetition_key(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", (text or "").lower())
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    normalized = re.sub(r"[^a-z0-9]+", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def repeated_clause_keep_limit(key: str) -> int:
    words = key.split()
    if len(words) <= 1 and key in {"si", "no", "ya", "ah"}:
        return 2
    return 1


def last_clause_key(text: str) -> str:
    parts = [part.strip() for part in (text or "").split(",") if part.strip()]
    return normalize_repetition_key(parts[-1]) if parts else normalize_repetition_key(text)


def collapse_repeated_clauses(text: str) -> tuple[str, dict[str, Any]]:
    original = (text or "").strip()
    parts = [part.strip() for part in original.split(",")]
    if len([part for part in parts if part]) < 3:
        return original, {}

    kept: list[str] = []
    removed: Counter[str] = Counter()
    previous_key = ""
    repeat_count = 0
    for part in parts:
        if not part:
            continue
        key = normalize_repetition_key(part)
        if not key:
            kept.append(part)
            previous_key = ""
            repeat_count = 0
            continue
        if key == previous_key:
            repeat_count += 1
        else:
            previous_key = key
            repeat_count = 1
        if repeat_count > repeated_clause_keep_limit(key):
            removed[key] += 1
            continue
        kept.append(part)

    if not removed:
        return original, {}
    cleaned = ", ".join(kept).strip()
    if cleaned and original and original[-1] in ".?!" and cleaned[-1] not in ".?!":
        cleaned += original[-1]
    return cleaned, {"removed_clauses": dict(removed)}


def internal_loop_candidates(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for segment in segments:
        text = (segment.get("text") or "").strip()
        cleaned, cleanup = collapse_repeated_clauses(text)
        removed = cleanup.get("removed_clauses") or {}
        if not removed:
            continue
        total_removed = sum(int(value) for value in removed.values())
        if total_removed < 3:
            continue
        example, count = max(removed.items(), key=lambda item: item[1])
        candidates.append(
            {
                "text": example,
                "count": int(count) + repeated_clause_keep_limit(example),
                "removed": total_removed,
                "internal": True,
                "start": round(float(segment.get("start", 0)), 3),
                "end": round(float(segment.get("end", 0)), 3),
                "cleaned": cleaned,
            }
        )
    return candidates


def merge_cleanup_data(*items: dict[str, Any]) -> dict[str, Any]:
    merged_dropped: Counter[str] = Counter()
    ranges: list[dict[str, Any]] = []
    text_cleanups: list[dict[str, Any]] = []
    for item in items:
        if not item:
            continue
        merged_dropped.update(item.get("dropped") or {})
        ranges.extend(item.get("ranges") or [])
        text_cleanups.extend(item.get("text_cleanups") or [])
    result: dict[str, Any] = {}
    if merged_dropped:
        result["dropped"] = dict(merged_dropped)
    if ranges:
        result["ranges"] = ranges
    if text_cleanups:
        result["text_cleanups"] = text_cleanups
    return result


def sanitize_internal_loop_segments(segments: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    cleaned_segments: list[dict[str, Any]] = []
    dropped: Counter[str] = Counter()
    ranges: list[dict[str, Any]] = []
    text_cleanups: list[dict[str, Any]] = []

    for segment in segments:
        start = float(segment.get("start", 0))
        end = float(segment.get("end", start))
        text = (segment.get("text") or "").strip()
        cleaned_text, text_cleanup = collapse_repeated_clauses(text)
        removed_clauses = text_cleanup.get("removed_clauses") or {}
        heavy_internal_loop = sum(int(value) for value in removed_clauses.values()) >= 3
        key = normalize_repetition_key(cleaned_text)
        duration = max(0.0, end - start)

        previous = cleaned_segments[-1] if cleaned_segments else None
        previous_key = normalize_repetition_key(previous.get("text", "")) if previous else ""
        previous_last_clause = last_clause_key(previous.get("text", "")) if previous else ""
        same_timestamp = (
            previous is not None
            and abs(start - float(previous.get("start", 0))) <= 0.15
            and abs(end - float(previous.get("end", 0))) <= 0.15
        )
        repeated_zero_segment = (
            duration <= 0.15
            and previous is not None
            and same_timestamp
            and key
            and (
                key == previous_key
                or key in previous_key
                or previous_key in key
                or (previous_last_clause and key.startswith(previous_last_clause))
            )
        )

        if (duration <= 0.15 and heavy_internal_loop) or repeated_zero_segment:
            dropped[key or normalize_repetition_key(text) or "segmento vacio"] += 1
            ranges.append(
                {
                    "start": round(start, 3),
                    "end": round(end, 3),
                    "text": normalize_repetition_key(text)[:120],
                    "reason": "internal_loop" if heavy_internal_loop else "zero_duration_repeat",
                }
            )
            continue

        updated = {**segment, "text": cleaned_text}
        if removed_clauses and cleaned_text != text:
            text_cleanups.append(
                {
                    "start": round(start, 3),
                    "end": round(end, 3),
                    "removed": dict(removed_clauses),
                    "before": text[:240],
                    "after": cleaned_text[:240],
                }
            )
        cleaned_segments.append(updated)

    return cleaned_segments, merge_cleanup_data(
        {"dropped": dict(dropped), "ranges": ranges},
        {"text_cleanups": text_cleanups},
    )


def convert_audio(project: dict[str, Any], tools: ToolPaths) -> Path:
    if not tools.ffmpeg:
        raise RuntimeError("No encontre ffmpeg. Ejecuta el setup otra vez.")
    ensure_not_stopped(project["id"])
    source = project_media_path(project, "source_path", "source")
    if not source or not source.is_file():
        raise RuntimeError("Audio original no encontrado dentro del proyecto.")
    wav_path = project_dir(project["id"]) / "audio_16k.wav"
    command = [
        tools.ffmpeg,
        "-y",
        "-i",
        str(source),
        "-vn",
        "-ar",
        "16000",
        "-ac",
        "1",
        "-c:a",
        "pcm_s16le",
        str(wav_path),
    ]
    result = run_command(project["id"], command)
    if result.returncode != 0 or not wav_path.exists():
        raise RuntimeError(f"ffmpeg fallo:\n{result.stdout[-2000:]}")
    return wav_path


def wav_duration(path: Path) -> float:
    with wave.open(str(path), "rb") as wav:
        frames = wav.getnframes()
        rate = wav.getframerate()
    return frames / rate if rate else 0.0


def format_seconds_label(seconds: float) -> str:
    seconds = max(0.0, float(seconds))
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    whole = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{whole:02d}"


def wav_rms(path: Path, start_seconds: float, end_seconds: float) -> float:
    if end_seconds <= start_seconds:
        return 0.0
    with wave.open(str(path), "rb") as wav:
        sample_width = wav.getsampwidth()
        channels = wav.getnchannels()
        rate = wav.getframerate()
        total_frames = wav.getnframes()
        if sample_width != 2 or rate <= 0:
            return 0.0
        start_frame = max(0, min(total_frames, int(round(start_seconds * rate))))
        end_frame = max(start_frame, min(total_frames, int(round(end_seconds * rate))))
        frame_count = end_frame - start_frame
        if frame_count <= 0:
            return 0.0
        wav.setpos(start_frame)
        raw = wav.readframes(frame_count)
    if not raw:
        return 0.0
    import array

    samples = array.array("h")
    samples.frombytes(raw)
    if sys.byteorder != "little":
        samples.byteswap()
    if channels > 1:
        samples = array.array("h", samples[::channels])
    if not samples:
        return 0.0
    square_sum = sum(float(sample) * float(sample) for sample in samples)
    return math.sqrt(square_sum / len(samples))


def whisper_supported_options(whisper_path: str) -> set[str]:
    cached = WHISPER_OPTION_CACHE.get(whisper_path)
    if cached is not None:
        return cached
    try:
        result = subprocess.run(
            [whisper_path, "--help"],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=8,
            check=False,
        )
        text = result.stdout or ""
    except Exception:
        text = ""
    options = {
        option
        for option in ("--max-context", "--no-fallback", "--suppress-nst", "--max-len", "--split-on-word")
        if option in text
    }
    WHISPER_OPTION_CACHE[whisper_path] = options
    return options


def whisper_strict_args(whisper_path: str) -> list[str]:
    supported = whisper_supported_options(whisper_path)
    args: list[str] = []
    if "--max-context" in supported:
        args.extend(["--max-context", "0"])
    if "--no-fallback" in supported:
        args.append("--no-fallback")
    if "--suppress-nst" in supported:
        args.append("--suppress-nst")
    if "--max-len" in supported:
        args.extend(["--max-len", str(WHISPER_MAX_LEN)])
    if "--split-on-word" in supported:
        args.append("--split-on-word")
    return args


def write_wav_chunk(source: Path, target: Path, start_seconds: float, duration_seconds: float) -> None:
    with wave.open(str(source), "rb") as src:
        params = src.getparams()
        rate = src.getframerate()
        start_frame = max(0, int(round(start_seconds * rate)))
        frame_count = max(0, int(round(duration_seconds * rate)))
        src.setpos(min(start_frame, src.getnframes()))
        frames = src.readframes(min(frame_count, src.getnframes() - min(start_frame, src.getnframes())))
    target.parent.mkdir(parents=True, exist_ok=True)
    with wave.open(str(target), "wb") as out:
        out.setparams(params)
        out.writeframes(frames)


def shift_segments(segments: list[dict[str, Any]], offset: float, id_prefix: str) -> list[dict[str, Any]]:
    shifted = []
    for idx, segment in enumerate(segments):
        start = round(float(segment.get("start", 0)) + offset, 3)
        end = round(float(segment.get("end", start)) + offset, 3)
        shifted.append(
            {
                **segment,
                "id": f"{id_prefix}-{idx:05d}",
                "start": start,
                "end": max(end, start),
            }
        )
    return shifted


def write_combined_whisper_json(output_base: Path, segments: list[dict[str, Any]]) -> None:
    payload = {
        "transcription": [
            {
                "offsets": {
                    "from": int(round(float(segment.get("start", 0)) * 1000)),
                    "to": int(round(float(segment.get("end", 0)) * 1000)),
                },
                "text": segment.get("text", ""),
            }
            for segment in segments
        ]
    }
    output_base.with_suffix(".json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    output_base.with_suffix(".txt").write_text("\n".join(segment.get("text", "") for segment in segments), encoding="utf-8")


def run_whisper_file(
    project: dict[str, Any],
    audio_path: Path,
    tools: ToolPaths,
    model_path: Path,
    output_base: Path,
    *,
    strict: bool = True,
    allow_empty: bool = False,
    max_threads: int = 4,
) -> list[dict[str, Any]]:
    if not tools.whisper:
        raise RuntimeError("No encontre whisper.cpp. Ejecuta el setup otra vez.")
    threads = str(max(1, min((os.cpu_count() or 2), max_threads)))
    command = [
        tools.whisper,
        "-m",
        str(model_path),
        "-f",
        str(audio_path),
        "-l",
        "es",
        "-t",
        threads,
        *(whisper_strict_args(tools.whisper) if strict else []),
        "-oj",
        "-otxt",
        "-osrt",
        "-ovtt",
        "-of",
        str(output_base),
    ]
    result = run_command(project["id"], command)
    ensure_not_stopped(project["id"])
    json_path = output_base.with_suffix(".json")
    if result.returncode != 0:
        raise RuntimeError(f"whisper.cpp fallo:\n{result.stdout[-2500:]}")
    if json_path.exists():
        segments = parse_whisper_json(json_path)
    else:
        segments = parse_whisper_stdout(result.stdout)
    if not segments:
        txt_path = output_base.with_suffix(".txt")
        if txt_path.exists():
            text = txt_path.read_text(encoding="utf-8", errors="ignore").strip()
            if text:
                segments = [{"id": "seg-00000", "start": 0, "end": 0, "speaker": "SPEAKER_00", "text": text}]
    if not segments:
        if allow_empty:
            return []
        raise RuntimeError("Whisper termino, pero no produjo texto interpretable.")
    return segments


def segment_word_count(segments: list[dict[str, Any]]) -> int:
    return sum(len((segment.get("text") or "").split()) for segment in segments)


def trim_segments_to_core(
    segments: list[dict[str, Any]],
    core_start: float,
    core_end: float,
) -> list[dict[str, Any]]:
    trimmed: list[dict[str, Any]] = []
    for segment in segments:
        start = float(segment.get("start", 0))
        end = float(segment.get("end", start))
        text = (segment.get("text") or "").strip()
        if not text:
            continue
        midpoint = (start + end) / 2 if end > start else start
        if midpoint < core_start or midpoint >= core_end:
            continue
        clipped_start = max(start, core_start)
        clipped_end = min(max(end, clipped_start), core_end)
        trimmed.append(
            {
                **segment,
                "start": round(clipped_start, 3),
                "end": round(clipped_end, 3),
                "text": text,
            }
        )
    return trimmed


def split_time_range(start: float, end: float, chunk_seconds: float) -> list[tuple[float, float]]:
    ranges = []
    cursor = start
    while cursor < end - 0.01:
        next_end = min(end, cursor + chunk_seconds)
        ranges.append((round(cursor, 3), round(next_end, 3)))
        cursor = next_end
    return ranges


def loop_candidates(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    keys = [normalize_repetition_key(segment.get("text", "")) for segment in segments]
    keys = [key for key in keys if key]
    candidates = internal_loop_candidates(segments)
    if len(keys) < 4:
        return candidates
    counts = Counter(keys)
    total = len(keys)
    for key, count in counts.items():
        words = key.split()
        ratio = count / total
        if len(words) == 1:
            suspicious = count >= 18 and ratio >= 0.45
        else:
            suspicious = len(words) <= 10 and count >= 6 and (ratio >= 0.25 or count >= 10)
        if suspicious:
            candidates.append({"text": key, "count": count, "ratio": round(ratio, 3)})

    tail = keys[-10:]
    if len(tail) >= 5:
        tail_counts = Counter(tail)
        key, count = tail_counts.most_common(1)[0]
        if count >= 4 and len(key.split()) <= 10 and not any(item["text"] == key for item in candidates):
            candidates.append({"text": key, "count": count, "ratio": round(count / len(tail), 3), "tail": True})
    return dedupe_loop_candidates(candidates)


def dedupe_loop_candidates(candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: dict[str, dict[str, Any]] = {}
    for item in candidates:
        key = normalize_repetition_key(str(item.get("text") or ""))
        if not key:
            continue
        count = int(item.get("count") or 0)
        current = merged.get(key)
        if current is None or count > int(current.get("count") or 0):
            merged[key] = {**item, "text": key, "count": count}
            continue
        for flag in ("internal", "tail"):
            if item.get(flag):
                current[flag] = True
        if item.get("removed"):
            current["removed"] = int(current.get("removed") or 0) + int(item.get("removed") or 0)
    return sorted(merged.values(), key=lambda item: int(item["count"]), reverse=True)


def non_silent_gaps(
    wav_path: Path,
    core_start: float,
    core_end: float,
    segments: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    gaps: list[dict[str, Any]] = []
    intervals = sorted(
        (
            max(core_start, float(segment.get("start", 0))),
            min(core_end, float(segment.get("end", 0))),
        )
        for segment in segments
        if (segment.get("text") or "").strip()
    )
    cursor = core_start
    for start, end in intervals:
        if start - cursor >= WHISPER_MAX_NON_SILENT_GAP_SECONDS:
            rms = wav_rms(wav_path, cursor, start)
            if rms >= WHISPER_RMS_SPEECH_THRESHOLD:
                gaps.append(
                    {
                        "start": round(cursor, 3),
                        "end": round(start, 3),
                        "seconds": round(start - cursor, 1),
                        "rms": round(rms, 1),
                    }
                )
        cursor = max(cursor, end)
    if core_end - cursor >= WHISPER_MAX_NON_SILENT_GAP_SECONDS:
        rms = wav_rms(wav_path, cursor, core_end)
        if rms >= WHISPER_RMS_SPEECH_THRESHOLD:
            gaps.append(
                {
                    "start": round(cursor, 3),
                    "end": round(core_end, 3),
                    "seconds": round(core_end - cursor, 1),
                    "rms": round(rms, 1),
                }
            )
    return gaps


def validate_whisper_chunk(
    wav_path: Path,
    segments: list[dict[str, Any]],
    core_start: float,
    core_end: float,
) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    rms = wav_rms(wav_path, core_start, core_end)
    if not segments and rms >= WHISPER_RMS_SPEECH_THRESHOLD:
        issues.append(
            {
                "type": "no_text_non_silent",
                "message": "Whisper no produjo texto en un tramo con audio detectable",
                "rms": round(rms, 1),
            }
        )

    loops = loop_candidates(segments)
    if loops:
        examples = ", ".join(f"{item['text']} ({item['count']})" for item in loops[:3])
        issues.append(
            {
                "type": "loop",
                "message": f"posible repeticion de Whisper: {examples}",
                "loops": loops[:6],
            }
        )

    gaps = non_silent_gaps(wav_path, core_start, core_end, segments)
    if gaps:
        first = gaps[0]
        issues.append(
            {
                "type": "non_silent_gap",
                "message": (
                    "gap con audio detectable "
                    f"{format_seconds_label(first['start'])}-{format_seconds_label(first['end'])}"
                ),
                "gaps": gaps[:8],
            }
        )

    return {
        "ok": not issues,
        "issues": issues,
        "segment_count": len(segments),
        "word_count": segment_word_count(segments),
        "rms": round(rms, 1),
    }


def summarize_quality_issues(issues: list[dict[str, Any]]) -> str:
    if not issues:
        return "resultado dudoso"
    return "; ".join(str(issue.get("message") or issue.get("type") or "resultado dudoso") for issue in issues[:3])


def sanitize_local_loop_segments(segments: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    segments, internal_cleanup = sanitize_internal_loop_segments(segments)
    if len(segments) < 8:
        return segments, internal_cleanup
    keys = [normalize_repetition_key(segment.get("text", "")) for segment in segments]
    counts = Counter(key for key in keys if key)
    total = max(1, len(keys))
    loop_keys = {
        key
        for key, count in counts.items()
        if key
        and 2 <= len(key.split()) <= 10
        and count >= 6
        and (count / total >= 0.25 or count >= 10)
    }
    if not loop_keys:
        return segments, internal_cleanup

    seen: Counter[str] = Counter()
    dropped: Counter[str] = Counter()
    dropped_ranges: list[dict[str, Any]] = []
    cleaned: list[dict[str, Any]] = []
    for segment, key in zip(segments, keys):
        if key in loop_keys:
            seen[key] += 1
            if seen[key] > 3:
                dropped[key] += 1
                dropped_ranges.append(
                    {
                        "start": round(float(segment.get("start", 0)), 3),
                        "end": round(float(segment.get("end", 0)), 3),
                        "text": key,
                    }
                )
                continue
        cleaned.append(segment)
    return cleaned, merge_cleanup_data(internal_cleanup, {"dropped": dict(dropped), "ranges": dropped_ranges})


def sanitize_whisper_loops_with_ranges(segments: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    segments, internal_cleanup = sanitize_internal_loop_segments(segments)
    if len(segments) < 80:
        return segments, internal_cleanup

    keys = [normalize_repetition_key(segment.get("text", "")) for segment in segments]
    counts = Counter(key for key in keys if key)
    threshold = max(30, int(len(segments) * 0.02))
    loop_keys = {
        key
        for key, count in counts.items()
        if count >= threshold and 2 <= len(key.split()) <= 10 and len(key) <= 100
    }
    if not loop_keys:
        return segments, internal_cleanup

    seen: Counter[str] = Counter()
    dropped: Counter[str] = Counter()
    ranges: list[dict[str, Any]] = []
    cleaned: list[dict[str, Any]] = []
    for segment, key in zip(segments, keys):
        if key in loop_keys:
            seen[key] += 1
            if seen[key] > 8:
                dropped[key] += 1
                ranges.append(
                    {
                        "start": round(float(segment.get("start", 0)), 3),
                        "end": round(float(segment.get("end", 0)), 3),
                        "text": key,
                    }
                )
                continue
        cleaned.append(segment)
    return cleaned, merge_cleanup_data(internal_cleanup, {"dropped": dict(dropped), "ranges": ranges})


def format_cleanup_warning(cleanup: dict[str, Any]) -> str:
    dropped = cleanup.get("dropped") or {}
    text_cleanups = cleanup.get("text_cleanups") or []
    if not dropped and not text_cleanups:
        return ""
    parts = []
    if dropped:
        total = sum(int(value) for value in dropped.values())
        examples = ", ".join(
            f"{key}: {count}" for key, count in sorted(dropped.items(), key=lambda item: item[1], reverse=True)[:3]
        )
        parts.append(f"removio {total} segmento(s) repetidos ({examples})")
    if text_cleanups:
        parts.append(f"corrigio repeticiones internas en {len(text_cleanups)} segmento(s)")
    ranges = cleanup.get("ranges") or []
    if ranges:
        first = ranges[0]
        span = f"{format_seconds_label(first['start'])}-{format_seconds_label(first['end'])}"
    elif text_cleanups:
        first = text_cleanups[0]
        span = f"{format_seconds_label(first['start'])}-{format_seconds_label(first['end'])}"
    else:
        span = "sin rango"
    return f"Whisper corrigio repeticiones: {' y '.join(parts)}. Primer rango para revisar: {span}."


def review_segments_for_quality(
    issues: list[dict[str, Any]],
    core_start: float,
    core_end: float,
) -> list[dict[str, Any]]:
    review_ranges: list[tuple[float, float, str]] = []
    for issue in issues:
        issue_type = issue.get("type")
        if issue_type == "no_text_non_silent":
            review_ranges.append((core_start, core_end, "audio_no_recuperado"))
        if issue_type == "non_silent_gap":
            for gap in issue.get("gaps") or []:
                review_ranges.append(
                    (
                        max(core_start, float(gap.get("start", core_start))),
                        min(core_end, float(gap.get("end", core_end))),
                        "gap_con_audio",
                    )
                )

    segments = []
    seen: set[tuple[int, int]] = set()
    for start, end, reason in review_ranges:
        if end - start < 0.5:
            continue
        key = (int(round(start * 1000)), int(round(end * 1000)))
        if key in seen:
            continue
        seen.add(key)
        segments.append(
            {
                "id": f"review-{key[0]}-{key[1]}",
                "start": round(start, 3),
                "end": round(end, 3),
                "speaker": "SPEAKER_00",
                "text": (
                    "[Audio no recuperado con confianza entre "
                    f"{format_seconds_label(start)} y {format_seconds_label(end)}. Revisar manualmente.]"
                ),
                "needs_review": True,
                "review_reason": reason,
            }
        )
    return segments


def run_whisper_chunked(
    project: dict[str, Any],
    wav_path: Path,
    tools: ToolPaths,
    model_path: Path,
    profile: str,
    *,
    reset_manifest: bool = False,
) -> list[dict[str, Any]]:
    project_id = project["id"]
    duration = wav_duration(wav_path)
    settings = processing_profile_settings(profile)
    base_chunk_seconds = max(20.0, float(settings["chunk_seconds"]))
    configured_retries = settings.get("retry_chunk_seconds") or WHISPER_RETRY_CHUNK_SECONDS
    retry_sizes = [float(size) for size in configured_retries if 5.0 <= float(size) < base_chunk_seconds]
    chunk_sizes = [base_chunk_seconds, *retry_sizes]
    overlap_seconds = max(0.0, min(float(settings["overlap_seconds"]), base_chunk_seconds / 4))
    max_threads = int(settings.get("max_threads") or 4)
    base_ranges = split_time_range(0.0, duration, base_chunk_seconds)
    total_chunks = max(1, len(base_ranges))
    chunk_dir = project_dir(project_id) / "whisper_chunks"
    output_base = project_dir(project_id) / "whisper"
    manifest_file = processing_manifest_path(project_id)
    audio_info = audio_signature(wav_path, duration)
    all_segments: list[dict[str, Any]] = []
    profile_name = str(settings["name"])

    def new_manifest() -> dict[str, Any]:
        return {
            "version": 1,
            "project_id": project_id,
            "created_at": now_ms(),
            "updated_at": now_ms(),
            "stage": "whisper",
            "model": model_path.name,
            "profile": profile_name,
            "audio": audio_info,
            "whisper": {
                "settings": {
                    "chunk_seconds": base_chunk_seconds,
                    "overlap_seconds": overlap_seconds,
                    "retry_chunk_seconds": retry_sizes,
                    "max_non_silent_gap_seconds": WHISPER_MAX_NON_SILENT_GAP_SECONDS,
                    "rms_speech_threshold": WHISPER_RMS_SPEECH_THRESHOLD,
                    "strict_args": whisper_strict_args(tools.whisper) if tools.whisper else [],
                    "max_threads": max_threads,
                },
                "chunks": [
                    {
                        "index": index,
                        "core_start": start,
                        "core_end": end,
                        "status": "pending",
                    }
                    for index, (start, end) in enumerate(base_ranges)
                ],
                "attempts": [],
                "warnings": [],
            },
        }

    def is_compatible_manifest(payload: Any) -> bool:
        if not isinstance(payload, dict):
            return False
        whisper = payload.get("whisper") or {}
        manifest_audio = payload.get("audio") or {}
        manifest_settings = whisper.get("settings") or {}
        return (
            payload.get("version") == 1
            and payload.get("model") == model_path.name
            and payload.get("profile") == profile_name
            and manifest_audio.get("sha256") == audio_info["sha256"]
            and abs(float(manifest_audio.get("duration", 0)) - duration) < 0.01
            and float(manifest_settings.get("chunk_seconds", 0)) == base_chunk_seconds
            and float(manifest_settings.get("overlap_seconds", 0)) == overlap_seconds
        )

    if reset_manifest:
        shutil.rmtree(chunk_dir, ignore_errors=True)
        manifest_file.unlink(missing_ok=True)

    existing_manifest = read_json_or_none(manifest_file)
    if is_compatible_manifest(existing_manifest):
        manifest = existing_manifest
        manifest["stage"] = "whisper"
        manifest["updated_at"] = now_ms()
    else:
        if existing_manifest:
            append_project_log(project_id, "Manifest de procesamiento incompatible; se reiniciaran checkpoints de Whisper.")
        shutil.rmtree(chunk_dir, ignore_errors=True)
        manifest = new_manifest()

    chunk_dir.mkdir(parents=True, exist_ok=True)

    def write_manifest() -> None:
        manifest["updated_at"] = now_ms()
        write_json_atomic(manifest_file, manifest)

    write_manifest()
    whisper_manifest = manifest["whisper"]
    attempts: list[dict[str, Any]] = list(whisper_manifest.get("attempts") or [])
    quality_warnings: list[str] = list(dict.fromkeys(whisper_manifest.get("warnings") or []))
    attempt_counter = max([int(attempt.get("index") or 0) for attempt in attempts] or [0])

    def accepted_chunk_path(chunk_index: int) -> Path:
        return chunk_dir / f"chunk_{chunk_index:04d}" / "accepted.json"

    def chunk_meta(chunk_index: int) -> dict[str, Any]:
        chunks = whisper_manifest.setdefault("chunks", [])
        while len(chunks) <= chunk_index:
            start, end = base_ranges[len(chunks)]
            chunks.append({"index": len(chunks), "core_start": start, "core_end": end, "status": "pending"})
        return chunks[chunk_index]

    def load_accepted_chunk(chunk_index: int) -> Optional[list[dict[str, Any]]]:
        meta = chunk_meta(chunk_index)
        if meta.get("status") != "accepted":
            return None
        payload = read_json_or_none(accepted_chunk_path(chunk_index))
        if not isinstance(payload, dict):
            meta["status"] = "pending"
            return None
        segments = payload.get("segments")
        if not isinstance(segments, list):
            meta["status"] = "pending"
            return None
        return segments

    def save_accepted_chunk(chunk_index: int, start: float, end: float, segments: list[dict[str, Any]]) -> None:
        path = accepted_chunk_path(chunk_index)
        payload = {
            "version": 1,
            "accepted_at": now_ms(),
            "index": chunk_index,
            "core_start": round(start, 3),
            "core_end": round(end, 3),
            "segments": segments,
        }
        write_json_atomic(path, payload)
        meta = chunk_meta(chunk_index)
        meta.update(
            {
                "status": "accepted",
                "accepted_path": str(path.relative_to(project_dir(project_id))),
                "segment_count": len(segments),
                "word_count": segment_word_count(segments),
                "updated_at": now_ms(),
            }
        )
        write_manifest()

    def save_partial_project() -> None:
        partial = [dict(segment) for segment in all_segments]
        for idx, segment in enumerate(partial):
            segment["id"] = f"seg-{idx:05d}"
        project["segments"] = partial
        project["speaker_labels"] = default_speaker_labels(partial, project.get("speaker_labels") or {})
        project["updated_at"] = now_ms()
        save_project(project)

    append_project_log(
        project_id,
        (
            "Whisper adaptativo: "
            f"duracion={duration:.1f}s chunk={base_chunk_seconds:.0f}s "
            f"overlap={overlap_seconds:.0f}s retries={','.join(str(int(size)) for size in retry_sizes) or 'none'} "
            f"profile={profile_name} threads={max_threads}"
        ),
    )

    def run_attempt(core_start: float, core_end: float, chunk_seconds: float, depth: int) -> tuple[list[dict[str, Any]], dict[str, Any]]:
        nonlocal attempt_counter
        ensure_not_stopped(project_id)
        attempt_counter += 1
        read_start = max(0.0, core_start - overlap_seconds)
        read_end = min(duration, core_end + overlap_seconds)
        chunk_path = chunk_dir / f"chunk_{attempt_counter:04d}_{int(core_start * 1000)}_{int(core_end * 1000)}.wav"
        chunk_output = chunk_dir / chunk_path.stem
        progress = round(25 + (core_start / max(duration, 1.0)) * 40)
        label = f"{format_seconds_label(core_start)}-{format_seconds_label(core_end)}"
        update_job(
            project_id,
            status="processing",
            step=f"Transcribiendo con Whisper {label}",
            progress=max(25, min(68, progress)),
            stage="whisper",
            retry_size=round(chunk_seconds, 1),
            can_resume=True,
        )
        append_project_log(
            project_id,
            (
                f"Whisper tramo {label}: intento chunk={chunk_seconds:.0f}s "
                f"read={format_seconds_label(read_start)}-{format_seconds_label(read_end)}"
            ),
        )
        try:
            write_wav_chunk(wav_path, chunk_path, read_start, read_end - read_start)
            raw_segments = run_whisper_file(
                project,
                chunk_path,
                tools,
                model_path,
                chunk_output,
                strict=True,
                allow_empty=True,
                max_threads=max_threads,
            )
            shifted = shift_segments(raw_segments, read_start, f"attempt-{attempt_counter:04d}")
            segments = trim_segments_to_core(shifted, core_start, core_end)
            quality = validate_whisper_chunk(wav_path, segments, core_start, core_end)
        except JobStopped:
            raise
        except Exception as exc:
            raw_error = str(exc)
            segments = []
            quality = {
                "ok": False,
                "fatal_error": is_fatal_whisper_error(raw_error),
                "issues": [{"type": "whisper_error", "message": compact_error_message(exc)}],
                "segment_count": 0,
                "word_count": 0,
                "rms": round(wav_rms(wav_path, core_start, core_end), 1),
            }
        finally:
            for path in [chunk_path, *[chunk_output.with_suffix(suffix) for suffix in (".json", ".txt", ".srt", ".vtt")]]:
                try:
                    path.unlink()
                except FileNotFoundError:
                    pass

        attempt = {
            "index": attempt_counter,
            "depth": depth,
            "chunk_seconds": round(chunk_seconds, 3),
            "core_start": round(core_start, 3),
            "core_end": round(core_end, 3),
            "read_start": round(read_start, 3),
            "read_end": round(read_end, 3),
            "accepted": False,
            **quality,
        }
        attempts.append(attempt)
        whisper_manifest["attempts"] = attempts
        write_manifest()
        return segments, attempt

    def transcribe_adaptive(core_start: float, core_end: float, size_index: int = 0, depth: int = 0) -> list[dict[str, Any]]:
        ensure_not_stopped(project_id)
        chunk_seconds = chunk_sizes[min(size_index, len(chunk_sizes) - 1)]
        segments, attempt = run_attempt(core_start, core_end, chunk_seconds, depth)
        if attempt.get("fatal_error"):
            raise RuntimeError(summarize_quality_issues(attempt.get("issues") or []))
        if attempt["ok"]:
            attempt["accepted"] = True
            return segments

        has_retry = size_index + 1 < len(chunk_sizes) and core_end - core_start > chunk_sizes[size_index + 1] + 5
        if has_retry:
            next_size = chunk_sizes[size_index + 1]
            append_project_log(
                project_id,
                (
                    f"Whisper reintentara {format_seconds_label(core_start)}-{format_seconds_label(core_end)} "
                    f"en tramos de {next_size:.0f}s: {summarize_quality_issues(attempt.get('issues') or [])}"
                ),
            )
            recovered: list[dict[str, Any]] = []
            for child_start, child_end in split_time_range(core_start, core_end, next_size):
                recovered.extend(transcribe_adaptive(child_start, child_end, size_index + 1, depth + 1))
            return recovered

        cleaned_segments, cleanup = sanitize_local_loop_segments(segments)
        if cleanup:
            attempt["local_cleanup"] = cleanup
            segments = cleaned_segments
            post_cleanup_quality = validate_whisper_chunk(wav_path, segments, core_start, core_end)
            attempt["post_cleanup_quality"] = post_cleanup_quality
            if post_cleanup_quality["ok"]:
                attempt["accepted"] = True
                attempt["accepted_after_cleanup"] = True
                cleanup_warning = format_cleanup_warning(cleanup)
                if cleanup_warning:
                    append_project_log(project_id, cleanup_warning)
                return segments
        attempt["accepted"] = True
        attempt["accepted_with_warning"] = True
        review_segments = review_segments_for_quality(attempt.get("issues") or [], core_start, core_end)
        if review_segments:
            segments = sorted([*segments, *review_segments], key=lambda item: (float(item.get("start", 0)), float(item.get("end", 0))))
        warning = (
            "Whisper requiere revision entre "
            f"{format_seconds_label(core_start)} y {format_seconds_label(core_end)}: "
            f"{summarize_quality_issues(attempt.get('issues') or [])}."
        )
        if cleanup:
            cleanup_warning = format_cleanup_warning(cleanup)
            if cleanup_warning:
                warning = f"{warning} {cleanup_warning}"
        if warning not in quality_warnings:
            quality_warnings.append(warning)
        whisper_manifest["warnings"] = quality_warnings
        update_job(project_id, status="processing", step=f"Revisar Whisper {format_seconds_label(core_start)}-{format_seconds_label(core_end)}", last_warning=warning)
        write_manifest()
        append_project_log(project_id, warning)
        return segments

    for chunk_index, (start, end) in enumerate(base_ranges):
        accepted_segments = load_accepted_chunk(chunk_index)
        if accepted_segments is not None:
            update_job(
                project_id,
                status="processing",
                step=f"Saltando tramo Whisper ya completado ({chunk_index + 1}/{total_chunks})",
                progress=round(25 + (chunk_index / total_chunks) * 40),
                stage="whisper",
                chunk_index=chunk_index + 1,
                chunk_total=total_chunks,
                can_resume=True,
            )
            all_segments.extend(accepted_segments)
            continue

        meta = chunk_meta(chunk_index)
        meta["status"] = "processing"
        meta["updated_at"] = now_ms()
        write_manifest()
        update_job(
            project_id,
            status="processing",
            step=f"Transcribiendo con Whisper ({chunk_index + 1}/{total_chunks})",
            progress=round(25 + (chunk_index / total_chunks) * 40),
            stage="whisper",
            chunk_index=chunk_index + 1,
            chunk_total=total_chunks,
            can_resume=True,
        )
        chunk_segments = transcribe_adaptive(start, end)
        save_accepted_chunk(chunk_index, start, end, chunk_segments)
        all_segments.extend(chunk_segments)
        save_partial_project()

    for idx, segment in enumerate(all_segments):
        segment["id"] = f"seg-{idx:05d}"
    write_combined_whisper_json(output_base, all_segments)
    quality_payload = {
        "version": 1,
        "created_at": now_ms(),
        "duration": round(duration, 3),
        "settings": {
            "chunk_seconds": base_chunk_seconds,
            "overlap_seconds": overlap_seconds,
            "retry_chunk_seconds": retry_sizes,
            "max_non_silent_gap_seconds": WHISPER_MAX_NON_SILENT_GAP_SECONDS,
            "rms_speech_threshold": WHISPER_RMS_SPEECH_THRESHOLD,
            "strict_args": whisper_strict_args(tools.whisper) if tools.whisper else [],
            "profile": profile_name,
            "max_threads": max_threads,
        },
        "summary": {
            "base_chunks": total_chunks,
            "attempts": len(attempts),
            "accepted_with_warning": sum(1 for attempt in attempts if attempt.get("accepted_with_warning")),
            "warnings": len(quality_warnings),
        },
        "warnings": quality_warnings,
        "attempts": attempts,
    }
    write_json_atomic(project_dir(project_id) / "whisper_quality.json", quality_payload)
    manifest["stage"] = "whisper_done"
    whisper_manifest["warnings"] = quality_warnings
    write_manifest()
    if quality_warnings:
        project.setdefault("warnings", []).extend(quality_warnings)
    if not all_segments:
        detail = quality_warnings[0] if quality_warnings else "Whisper no produjo texto."
        raise RuntimeError(f"Whisper no produjo segmentos transcritos. {detail}")
    return all_segments


def run_whisper(
    project: dict[str, Any],
    wav_path: Path,
    tools: ToolPaths,
    model_name: str,
    profile: str,
    *,
    reset_manifest: bool = False,
) -> list[dict[str, Any]]:
    if not tools.whisper:
        raise RuntimeError("No encontre whisper.cpp. Ejecuta el setup otra vez.")
    ensure_not_stopped(project["id"])
    model_path = MODELS_DIR / "whisper" / model_name
    if not model_path.exists():
        models = available_models()
        if not models:
            raise RuntimeError("No hay modelos Whisper en models/whisper. Ejecuta el setup otra vez.")
        model_path = Path(models[0]["path"])
    integrity_error = model_integrity_error(model_path)
    if integrity_error:
        raise RuntimeError(
            f"Modelo Whisper invalido ({model_path.name}): {integrity_error}. "
            "Borra ese archivo y descargalo de nuevo."
        )

    segments = run_whisper_chunked(project, wav_path, tools, model_path, profile, reset_manifest=reset_manifest)

    segments, cleanup = sanitize_whisper_loops_with_ranges(segments)
    if cleanup:
        message = format_cleanup_warning(cleanup)
        append_project_log(project["id"], message)
        project.setdefault("warnings", []).append(message)
        quality_path = project_dir(project["id"]) / "whisper_quality.json"
        if quality_path.exists():
            try:
                quality_payload = json.loads(quality_path.read_text(encoding="utf-8"))
                quality_payload["final_cleanup"] = cleanup
                quality_payload.setdefault("warnings", []).append(message)
                quality_path.write_text(json.dumps(quality_payload, ensure_ascii=False, indent=2), encoding="utf-8")
            except (OSError, json.JSONDecodeError):
                append_project_log(project["id"], "No se pudo actualizar whisper_quality.json con cleanup final.")
        write_combined_whisper_json(project_dir(project["id"]) / "whisper", segments)
    return segments


def read_wav_float32(path: Path) -> tuple[Any, int]:
    import numpy as np

    with wave.open(str(path), "rb") as wav:
        sample_rate = wav.getframerate()
        channels = wav.getnchannels()
        sample_width = wav.getsampwidth()
        frames = wav.readframes(wav.getnframes())
    if sample_width != 2:
        raise RuntimeError("La diarizacion espera WAV PCM 16-bit.")
    samples = np.frombuffer(frames, dtype="<i2").astype("float32") / 32768.0
    if channels > 1:
        samples = samples.reshape(-1, channels).mean(axis=1)
    return samples, sample_rate


def run_diarization(wav_path: Path, output_path: Path, num_speakers: Optional[int] = None) -> list[dict[str, Any]]:
    project_id = output_path.parent.name
    ensure_not_stopped(project_id)
    command = [
        sys.executable,
        str(ROOT_DIR / "scripts" / "diarize_file.py"),
        "--wav",
        str(wav_path),
        "--models-dir",
        str(MODELS_DIR),
        "--output",
        str(output_path),
        "--chunk-seconds",
        str(DIARIZATION_CHUNK_SECONDS),
        "--chunk-overlap-seconds",
        str(DIARIZATION_CHUNK_OVERLAP_SECONDS),
    ]
    if num_speakers:
        command.extend(["--speakers", str(num_speakers)])
    append_project_log(project_id, "RUN " + " ".join(command))
    env = os.environ.copy()
    env["PYTHONUNBUFFERED"] = "1"
    process = subprocess.Popen(
        command,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        bufsize=1,
        env=env,
    )
    output_tail: list[str] = []
    with JOBS_LOCK:
        RUNNING_PROCESSES[project_id] = process
    try:
        if process.stdout:
            for raw_line in process.stdout:
                line = raw_line.rstrip()
                if not line:
                    continue
                output_tail.append(line)
                output_tail = output_tail[-120:]
                append_project_log(project_id, line)
                match = DIARIZE_PROGRESS_RE.match(line)
                if match:
                    percent = float(match.group("percent"))
                    message = match.group("message").strip() or "Separando hablantes"
                    mapped_progress = round(70 + (max(0.0, min(100.0, percent)) * 0.18))
                    update_job(
                        project_id,
                        status="processing",
                        step=message,
                        progress=mapped_progress,
                        stage="diarization",
                        can_resume=True,
                    )
        returncode = process.wait()
    finally:
        with JOBS_LOCK:
            if RUNNING_PROCESSES.get(project_id) is process:
                RUNNING_PROCESSES.pop(project_id, None)

    requested = get_stop_request(project_id)
    if requested in {"paused", "cancelled"}:
        append_project_log(project_id, f"STOP {requested}")
        raise JobStopped(requested)

    detail = "\n".join(output_tail).strip()
    if returncode != 0:
        if returncode < 0:
            detail = f"Proceso terminado por senal {-returncode}." + (f"\n{detail}" if detail else "")
        elif not detail:
            detail = f"Codigo de salida {returncode}, sin salida adicional."
        append_project_log(project_id, f"EXIT {returncode}\n{detail[-4000:]}")
        raise RuntimeError(f"sherpa-onnx fallo o fue terminado:\n{detail[-2000:]}")
    append_project_log(project_id, f"EXIT {returncode}")
    if not output_path.exists():
        raise RuntimeError("sherpa-onnx termino, pero no genero diarizacion.")
    return json.loads(output_path.read_text(encoding="utf-8"))


def overlap(a_start: float, a_end: float, b_start: float, b_end: float) -> float:
    return max(0.0, min(a_end, b_end) - max(a_start, b_start))


def turn_duration(turn: dict[str, Any]) -> float:
    return max(0.0, float(turn.get("end", 0)) - float(turn.get("start", 0)))


def segment_duration(segment: dict[str, Any]) -> float:
    return max(0.0, float(segment.get("end", 0)) - float(segment.get("start", 0)))


def segment_word_count_value(segment: dict[str, Any]) -> int:
    return len((segment.get("text") or "").split())


def auto_split_base_id(segment_id: Any) -> Optional[str]:
    match = re.match(r"^(.+)-sp\d+$", str(segment_id or ""))
    return match.group(1) if match else None


def merge_auto_split_segments(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    index = 0
    while index < len(segments):
        segment = segments[index]
        base_id = auto_split_base_id(segment.get("id"))
        if not base_id:
            merged.append(dict(segment))
            index += 1
            continue

        group = [segment]
        index += 1
        while index < len(segments) and auto_split_base_id(segments[index].get("id")) == base_id:
            group.append(segments[index])
            index += 1

        first = group[0]
        combined = dict(first)
        combined["id"] = base_id
        combined["start"] = round(min(float(item.get("start", 0)) for item in group), 3)
        combined["end"] = round(max(float(item.get("end", 0)) for item in group), 3)
        combined["text"] = " ".join((item.get("text") or "").strip() for item in group).replace("  ", " ").strip()
        merged.append(combined)
    return merged


def normalize_turns(turns: list[dict[str, Any]]) -> list[dict[str, Any]]:
    normalized = []
    for turn in turns:
        start = float(turn.get("start", 0))
        end = float(turn.get("end", start))
        if end - start <= 0.08:
            continue
        normalized.append(
            {
                "start": round(start, 3),
                "end": round(end, 3),
                "speaker": str(turn.get("speaker") or "SPEAKER_00"),
            }
        )
    return sorted(normalized, key=lambda item: (float(item["start"]), float(item["end"])))


def merge_nearby_same_speaker_turns(turns: list[dict[str, Any]], max_gap: float = 0.75) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    for turn in turns:
        item = dict(turn)
        if (
            merged
            and merged[-1]["speaker"] == item["speaker"]
            and float(item["start"]) - float(merged[-1]["end"]) <= max_gap
        ):
            merged[-1]["end"] = round(max(float(merged[-1]["end"]), float(item["end"])), 3)
            continue
        merged.append(item)
    return merged


def absorb_micro_turns(turns: list[dict[str, Any]], min_duration: float = 0.45) -> list[dict[str, Any]]:
    if not turns:
        return []
    cleaned: list[dict[str, Any]] = []
    for index, turn in enumerate(turns):
        if turn_duration(turn) >= min_duration:
            cleaned.append(dict(turn))
            continue
        previous_turn = cleaned[-1] if cleaned else None
        next_turn = turns[index + 1] if index + 1 < len(turns) else None
        if (
            previous_turn
            and next_turn
            and previous_turn["speaker"] == next_turn["speaker"]
            and float(turn["start"]) - float(previous_turn["end"]) <= 0.75
            and float(next_turn["start"]) - float(turn["end"]) <= 0.75
        ):
            previous_turn["end"] = round(max(float(previous_turn["end"]), float(next_turn["end"])), 3)
            continue
        if previous_turn and float(turn["start"]) - float(previous_turn["end"]) <= 0.35:
            continue
        if next_turn and float(next_turn["start"]) - float(turn["end"]) <= 0.35:
            continue
        cleaned.append(dict(turn))
    return merge_nearby_same_speaker_turns(cleaned)


def resolve_turn_overlaps(turns: list[dict[str, Any]]) -> list[dict[str, Any]]:
    resolved: list[dict[str, Any]] = []
    for turn in turns:
        current = dict(turn)
        if not resolved:
            resolved.append(current)
            continue
        previous = resolved[-1]
        if float(current["start"]) >= float(previous["end"]):
            resolved.append(current)
            continue
        if previous["speaker"] == current["speaker"]:
            previous["end"] = round(max(float(previous["end"]), float(current["end"])), 3)
            continue

        previous_duration = turn_duration(previous)
        current_duration = turn_duration(current)
        if current_duration <= 1.2 and previous_duration >= current_duration * 2:
            continue
        if previous_duration <= 1.2 and current_duration >= previous_duration * 2:
            resolved.pop()
            resolved.append(current)
            continue

        boundary = round((float(current["start"]) + float(previous["end"])) / 2, 3)
        previous["end"] = boundary
        current["start"] = boundary
        if turn_duration(previous) <= 0.08:
            resolved.pop()
        if turn_duration(current) > 0.08:
            resolved.append(current)
    return resolved


def suppress_short_turn_islands(turns: list[dict[str, Any]], max_duration: float = 1.2) -> list[dict[str, Any]]:
    if len(turns) < 3:
        return turns
    cleaned = [dict(turn) for turn in turns]
    for index in range(1, len(cleaned) - 1):
        previous_turn = cleaned[index - 1]
        current = cleaned[index]
        next_turn = cleaned[index + 1]
        if (
            current["speaker"] != previous_turn["speaker"]
            and previous_turn["speaker"] == next_turn["speaker"]
            and turn_duration(current) <= max_duration
            and float(current["start"]) - float(previous_turn["end"]) <= 0.75
            and float(next_turn["start"]) - float(current["end"]) <= 0.75
        ):
            current["speaker"] = previous_turn["speaker"]
    return merge_nearby_same_speaker_turns(cleaned)


def postprocess_diarization_turns(turns: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cleaned = normalize_turns(turns)
    cleaned = absorb_micro_turns(cleaned)
    cleaned = resolve_turn_overlaps(cleaned)
    cleaned = merge_nearby_same_speaker_turns(cleaned)
    cleaned = suppress_short_turn_islands(cleaned)
    return resolve_turn_overlaps(merge_nearby_same_speaker_turns(cleaned))


def best_speaker_for_segment(
    start: float,
    end: float,
    turns: list[dict[str, Any]],
    fallback: str = "SPEAKER_00",
) -> str:
    best_speaker = fallback
    best_overlap = 0.0
    midpoint = (start + end) / 2
    nearest_distance = float("inf")
    nearest_speaker = fallback
    for turn in turns:
        turn_start = float(turn["start"])
        turn_end = float(turn["end"])
        value = overlap(start, end, turn_start, turn_end)
        if value > best_overlap:
            best_overlap = value
            best_speaker = turn["speaker"]
        center = (turn_start + turn_end) / 2
        distance = abs(center - midpoint)
        if distance < nearest_distance:
            nearest_distance = distance
            nearest_speaker = turn["speaker"]
    return best_speaker if best_overlap > 0 else nearest_speaker


def speaker_parts_for_segment(
    start: float,
    end: float,
    turns: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    duration = max(0.0, end - start)
    if duration <= 0:
        return []
    min_overlap = 0.35 if duration >= 4 else 0.18
    parts = []
    for turn in turns:
        turn_start = float(turn["start"])
        turn_end = float(turn["end"])
        value = overlap(start, end, turn_start, turn_end)
        if value < min_overlap:
            continue
        parts.append(
            {
                "start": max(start, turn_start),
                "end": min(end, turn_end),
                "speaker": turn["speaker"],
            }
        )
    if not parts:
        return []
    parts.sort(key=lambda item: (item["start"], item["end"]))

    merged: list[dict[str, Any]] = []
    for part in parts:
        if not merged:
            merged.append(part)
            continue
        previous = merged[-1]
        if part["speaker"] == previous["speaker"] and part["start"] - previous["end"] <= 0.45:
            previous["end"] = max(previous["end"], part["end"])
        else:
            merged.append(part)

    return [part for part in merged if part["end"] - part["start"] >= 0.12]


def normalized_text_key(text: str) -> str:
    value = unicodedata.normalize("NFKD", text or "")
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^a-zA-Z0-9 ]+", " ", value).lower()
    return re.sub(r"\s+", " ", value).strip()


def is_question_like(text: str) -> bool:
    clean = normalized_text_key(text)
    if "?" in text or "¿" in text:
        return True
    starters = (
        "cuanto",
        "cuanta",
        "cuantos",
        "cuantas",
        "cual",
        "cuales",
        "como",
        "cuando",
        "donde",
        "quien",
        "quienes",
        "tiene",
        "tienen",
        "ustedes participan",
    )
    return any(clean.startswith(starter) for starter in starters)


def is_prompt_like(text: str) -> bool:
    clean = normalized_text_key(text)
    prompt_markers = (
        "ahora le voy a preguntar",
        "recapitulando",
        "integrante",
        "nivel educativo",
        "nacionalidad",
        "hogar",
        "dependencia funcional",
        "movilidad",
        "me cuenta",
        "le voy a preguntar",
        "haciamos",
        "tratamos",
        "tema es de tiempo",
        "mas agil",
        "para ustedes",
        "para nosotros",
        "gracias por recibirnos",
        "para el segundo",
        "para la segunda",
        "para el tercer",
        "para la tercera",
        "super",
    )
    return any(marker in clean for marker in prompt_markers)


def is_interviewer_ack_like(text: str) -> bool:
    clean = normalized_text_key(text)
    return clean in {
        "ya",
        "ah ya",
        "ya super",
        "super",
        "perfecto",
        "ok",
        "claro",
        "entiendo",
    }


def is_clear_short_response(text: str) -> bool:
    clean = normalized_text_key(text)
    if clean in {"si", "no", "ya", "claro", "exacto", "tambien"}:
        return True
    if re.fullmatch(r"(si|no)( si| no){0,3}", clean):
        return True
    if re.fullmatch(r"\d{1,3}", clean):
        return True
    return False


def role_speaker_ids(existing_labels: Optional[dict[str, str]] = None) -> tuple[str, str]:
    labels = existing_labels or {}
    interviewer = None
    interviewee = None
    for speaker, label in labels.items():
        clean = normalized_text_key(label)
        if "entrevistador" in clean:
            interviewer = speaker
        if "entrevistada" in clean or "entrevistado" in clean:
            interviewee = speaker
    return interviewer or "SPEAKER_00", interviewee or "SPEAKER_01"


def speaker_overlap_scores(
    start: float,
    end: float,
    turns: list[dict[str, Any]],
) -> dict[str, float]:
    scores: dict[str, float] = {}
    for turn in turns:
        value = overlap(start, end, float(turn["start"]), float(turn["end"]))
        if value > 0:
            scores[str(turn["speaker"])] = scores.get(str(turn["speaker"]), 0.0) + value
    return scores


def choose_interview_speaker(
    segment: dict[str, Any],
    turns: list[dict[str, Any]],
    existing_labels: Optional[dict[str, str]],
    previous_speaker: Optional[str],
    previous_end: Optional[float],
) -> str:
    start = float(segment.get("start", 0))
    end = float(segment.get("end", start))
    text = segment.get("text") or ""
    words = segment_word_count_value(segment)
    duration = max(0.0, end - start)
    interviewer, interviewee = role_speaker_ids(existing_labels)
    speakers = sorted({interviewer, interviewee, *(turn.get("speaker", "SPEAKER_00") for turn in turns)})
    scores = {speaker: 0.0 for speaker in speakers}

    acoustic_scores = speaker_overlap_scores(start, end, turns)
    for speaker, value in acoustic_scores.items():
        scores[speaker] = scores.get(speaker, 0.0) + value * 0.35
    if not acoustic_scores and turns:
        nearest = best_speaker_for_segment(start, end, turns, segment.get("speaker") or interviewer)
        scores[nearest] = scores.get(nearest, 0.0) + 0.4

    question = is_question_like(text)
    prompt = is_prompt_like(text)
    clear_response = is_clear_short_response(text)
    if question:
        scores[interviewer] = scores.get(interviewer, 0.0) + 4.0
    if prompt:
        scores[interviewer] = scores.get(interviewer, 0.0) + 3.0
    if is_interviewer_ack_like(text) and words <= 4:
        scores[interviewer] = scores.get(interviewer, 0.0) + 2.0
    if not question and not prompt:
        if previous_speaker == interviewer and (words >= 3 or clear_response):
            scores[interviewee] = scores.get(interviewee, 0.0) + 3.0
        if previous_speaker == interviewee and not is_interviewer_ack_like(text):
            scores[interviewee] = scores.get(interviewee, 0.0) + 1.0
        if words >= 8:
            scores[interviewee] = scores.get(interviewee, 0.0) + 2.0
        if clear_response:
            scores[interviewee] = scores.get(interviewee, 0.0) + 1.5

    if previous_speaker and duration <= 1.2 and not clear_response:
        scores[previous_speaker] = scores.get(previous_speaker, 0.0) + 1.4
    if previous_speaker and previous_end is not None and start - previous_end <= 0.75 and not question:
        scores[previous_speaker] = scores.get(previous_speaker, 0.0) + 0.4

    return max(scores.items(), key=lambda item: (item[1], item[0]))[0]


def should_split_segment_by_speaker(
    segment: dict[str, Any],
    parts: list[dict[str, Any]],
) -> bool:
    start = float(segment.get("start", 0))
    end = float(segment.get("end", start))
    text = segment.get("text") or ""
    if end - start < 6.0 or len(text.split()) < 18:
        return False
    if len(parts) < 2 or len(parts) > 4 or len({part["speaker"] for part in parts}) < 2:
        return False
    if any(float(part["end"]) - float(part["start"]) < 1.5 for part in parts):
        return False
    if min(float(parts[0]["end"]) - start, end - float(parts[-1]["start"])) < 1.5:
        return False
    return True


def split_text_by_durations(text: str, durations: list[float]) -> list[str]:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    count = len(durations)
    if count <= 1:
        return [cleaned]
    if not cleaned:
        return [""] * count

    words = cleaned.split(" ")
    if len(words) < count:
        return [cleaned] + [""] * (count - 1)

    total = sum(max(0.001, duration) for duration in durations)
    chunks: list[str] = []
    cursor = 0
    elapsed = 0.0
    for idx, duration in enumerate(durations[:-1]):
        elapsed += max(0.001, duration)
        target = round(len(words) * (elapsed / total))
        remaining_chunks = count - idx - 1
        min_cut = cursor + 1
        max_cut = len(words) - remaining_chunks
        cut = min(max(target, min_cut), max_cut)
        chunks.append(" ".join(words[cursor:cut]).strip())
        cursor = cut
    chunks.append(" ".join(words[cursor:]).strip())
    return chunks


def smooth_short_segment_islands(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if len(segments) < 3:
        return segments
    smoothed = [dict(segment) for segment in segments]
    for index in range(1, len(smoothed) - 1):
        previous = smoothed[index - 1]
        current = smoothed[index]
        next_segment = smoothed[index + 1]
        if (
            current.get("speaker") != previous.get("speaker")
            and previous.get("speaker") == next_segment.get("speaker")
            and segment_duration(current) <= 1.2
            and not is_clear_short_response(current.get("text") or "")
            and not is_question_like(current.get("text") or "")
            and not is_prompt_like(current.get("text") or "")
            and not is_interviewer_ack_like(current.get("text") or "")
        ):
            current["speaker"] = previous.get("speaker")
    return smoothed


def assign_speakers_from_processed_turns(
    segments: list[dict[str, Any]],
    turns: list[dict[str, Any]],
    existing_labels: Optional[dict[str, str]] = None,
) -> list[dict[str, Any]]:
    if not turns:
        return segments

    assigned: list[dict[str, Any]] = []
    previous_speaker: Optional[str] = None
    previous_end: Optional[float] = None
    for index, segment in enumerate(merge_auto_split_segments(segments)):
        start = float(segment.get("start", 0))
        end = float(segment.get("end", start))
        parts = speaker_parts_for_segment(start, end, turns)
        if should_split_segment_by_speaker(segment, parts):
            durations = [part["end"] - part["start"] for part in parts]
            text_parts = split_text_by_durations(segment.get("text") or "", durations)
            split_segments = []
            base_id = segment.get("id") or f"seg-{index}"
            for part_idx, (part, part_text) in enumerate(zip(parts, text_parts)):
                if not part_text:
                    continue
                item = dict(segment)
                item["id"] = f"{base_id}-sp{part_idx}"
                item["start"] = round(float(part["start"]), 3)
                item["end"] = round(float(part["end"]), 3)
                item["text"] = part_text
                item["speaker"] = part["speaker"]
                split_segments.append(item)
                previous_speaker = item["speaker"]
                previous_end = float(item["end"])
            if len(split_segments) >= 2:
                assigned.extend(split_segments)
                continue

        item = dict(segment)
        item["id"] = item.get("id") or f"seg-{index}"
        item["speaker"] = choose_interview_speaker(item, turns, existing_labels, previous_speaker, previous_end)
        assigned.append(item)
        previous_speaker = item["speaker"]
        previous_end = float(item.get("end", start))
    return smooth_short_segment_islands(assigned)


def assign_speakers(
    segments: list[dict[str, Any]],
    turns: list[dict[str, Any]],
    existing_labels: Optional[dict[str, str]] = None,
) -> list[dict[str, Any]]:
    if not turns:
        return segments
    processed_turns = postprocess_diarization_turns(turns)
    return assign_speakers_from_processed_turns(segments, processed_turns, existing_labels)


def default_speaker_labels(
    segments: list[dict[str, Any]],
    existing_labels: Optional[dict[str, str]] = None,
) -> dict[str, str]:
    speakers = sorted({segment.get("speaker", "SPEAKER_00") for segment in segments})
    defaults = ["Entrevistador/a", "Entrevistada/o", "Otra persona"]
    labels = {}
    for idx, speaker in enumerate(speakers):
        existing = (existing_labels or {}).get(speaker)
        labels[speaker] = existing or (defaults[idx] if idx < len(defaults) else speaker)
    return labels


def count_speaker_switches(items: list[dict[str, Any]]) -> int:
    return sum(1 for current, next_item in zip(items, items[1:]) if current.get("speaker") != next_item.get("speaker"))


def count_short_segment_islands(segments: list[dict[str, Any]], max_duration: float = 1.2) -> int:
    total = 0
    for index in range(1, len(segments) - 1):
        previous = segments[index - 1]
        current = segments[index]
        next_segment = segments[index + 1]
        if (
            current.get("speaker") != previous.get("speaker")
            and previous.get("speaker") == next_segment.get("speaker")
            and segment_duration(current) <= max_duration
            and not is_clear_short_response(current.get("text") or "")
        ):
            total += 1
    return total


def adjacent_overlap_stats(items: list[dict[str, Any]]) -> dict[str, Any]:
    count = 0
    seconds = 0.0
    for current, next_item in zip(items, items[1:]):
        if current.get("speaker") == next_item.get("speaker"):
            continue
        value = overlap(
            float(current.get("start", 0)),
            float(current.get("end", 0)),
            float(next_item.get("start", 0)),
            float(next_item.get("end", 0)),
        )
        if value > 0:
            count += 1
            seconds += value
    return {"count": count, "seconds": round(seconds, 3)}


def speaker_distribution(items: list[dict[str, Any]]) -> dict[str, dict[str, float | int]]:
    distribution: dict[str, dict[str, float | int]] = {}
    for item in items:
        speaker = str(item.get("speaker") or "SPEAKER_00")
        bucket = distribution.setdefault(speaker, {"count": 0, "seconds": 0.0})
        bucket["count"] = int(bucket["count"]) + 1
        bucket["seconds"] = round(float(bucket["seconds"]) + segment_duration(item), 3)
    return distribution


def compute_diarization_quality(
    raw_turns: list[dict[str, Any]],
    processed_turns: list[dict[str, Any]],
    segments: list[dict[str, Any]],
) -> dict[str, Any]:
    duration = max(
        [float(item.get("end", 0)) for item in raw_turns + processed_turns + segments] or [0.0]
    )
    minutes = max(duration / 60.0, 1.0)
    raw_overlap = adjacent_overlap_stats(normalize_turns(raw_turns))
    segment_overlap = adjacent_overlap_stats(segments)
    raw_micro_turns = sum(1 for turn in normalize_turns(raw_turns) if turn_duration(turn) < 0.45)
    short_islands = count_short_segment_islands(segments)
    split_segments = [segment for segment in segments if auto_split_base_id(segment.get("id"))]
    short_split_segments = [segment for segment in split_segments if segment_duration(segment) <= 2.0]
    segment_switches = count_speaker_switches(segments)
    processed_switches = count_speaker_switches(processed_turns)
    warnings = []
    if short_islands >= 20:
        warnings.append(f"{short_islands} islas cortas de hablante")
    if raw_overlap["count"] >= 20:
        warnings.append(f"{raw_overlap['count']} solapes crudos")
    if raw_micro_turns >= 20:
        warnings.append(f"{raw_micro_turns} micro-turnos")
    return {
        "version": 1,
        "created_at": now_ms(),
        "duration": round(duration, 3),
        "raw_turns": len(raw_turns),
        "processed_turns": len(processed_turns),
        "segments": len(segments),
        "raw_micro_turns": raw_micro_turns,
        "raw_adjacent_overlaps": raw_overlap,
        "segment_adjacent_overlaps": segment_overlap,
        "short_speaker_islands": short_islands,
        "auto_split_segments": len(split_segments),
        "short_auto_split_segments": len(short_split_segments),
        "processed_turn_switches_per_minute": round(processed_switches / minutes, 2),
        "segment_switches_per_minute": round(segment_switches / minutes, 2),
        "speaker_distribution": speaker_distribution(segments),
        "warnings": warnings,
    }


def format_diarization_quality_warning(quality: dict[str, Any]) -> str:
    short_islands = int(quality.get("short_speaker_islands") or 0)
    raw_overlap = quality.get("raw_adjacent_overlaps") or {}
    raw_overlap_count = int(raw_overlap.get("count") or 0)
    raw_micro_turns = int(quality.get("raw_micro_turns") or 0)
    short_auto_splits = int(quality.get("short_auto_split_segments") or 0)

    notes = []
    if raw_overlap_count >= 20 or raw_micro_turns >= 20:
        raw_parts = []
        if raw_overlap_count >= 20:
            raw_parts.append(f"{raw_overlap_count} solapes crudos")
        if raw_micro_turns >= 20:
            raw_parts.append(f"{raw_micro_turns} micro-turnos")
        notes.append(f"se suavizo una diarizacion inestable ({', '.join(raw_parts)})")
    if short_islands >= 20:
        notes.append(f"quedan {short_islands} cambios breves de hablante para revisar")
    if short_auto_splits >= 20:
        notes.append(f"quedan {short_auto_splits} segmentos divididos muy cortos")
    if not notes:
        return ""
    return "Separacion de hablantes: " + "; ".join(notes) + "."


def write_diarization_quality(project_id: str, quality: dict[str, Any]) -> None:
    (project_dir(project_id) / "diarization_quality.json").write_text(
        json.dumps(quality, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def relabel_segments_with_diagnostics(
    project_id: str,
    segments: list[dict[str, Any]],
    turns: list[dict[str, Any]],
    existing_labels: Optional[dict[str, str]] = None,
) -> tuple[list[dict[str, Any]], dict[str, str], dict[str, Any], str]:
    processed_turns = postprocess_diarization_turns(turns)
    assigned = assign_speakers_from_processed_turns(segments, processed_turns, existing_labels)
    labels = default_speaker_labels(assigned, existing_labels)
    quality = compute_diarization_quality(turns, processed_turns, assigned)
    write_diarization_quality(project_id, quality)
    warning = format_diarization_quality_warning(quality)
    return assigned, labels, quality, warning


def clear_diarization_warnings(warnings: list[str]) -> list[str]:
    cleaned = []
    for warning in warnings:
        text = warning.lower()
        if "diariz" in text or text.startswith("separacion de hablantes"):
            continue
        cleaned.append(warning)
    return cleaned


def compact_error_message(error: Exception) -> str:
    text = str(error).strip()
    if not text:
        return "error sin detalle."
    for line in text.splitlines():
        line = line.strip()
        lowered = line.lower()
        if not line:
            continue
        if lowered in {"whisper.cpp fallo:", "sherpa-onnx fallo o fue terminado:"}:
            continue
        if lowered.endswith("fallo:") and len(line) < 80:
            continue
        if line:
            return line[:300]
    return text[:300]


def parse_java_major_version(text: str) -> Optional[int]:
    match = re.search(r'version\s+"?([0-9]+)(?:\.([0-9]+))?', text or "", re.IGNORECASE)
    if not match:
        return None
    first = int(match.group(1))
    second = int(match.group(2) or 0)
    if first == 1 and second:
        return second
    return first


def java_install_hint() -> str:
    system = platform.system()
    if system == "Windows":
        return "Instala Java 17 o superior de 64 bits. Con winget: winget install -e --id EclipseAdoptium.Temurin.17.JRE"
    if system == "Darwin":
        return "Instala Java 17 o superior. Con Homebrew: brew install --cask temurin"
    return "Instala Java 17 o superior y vuelve a activar el corrector."


def java_candidate_paths() -> list[Path]:
    executable = "java.exe" if platform.system() == "Windows" else "java"
    candidates: list[Path] = []
    seen: set[str] = set()

    def add(candidate: Optional[Path | str]) -> None:
        if not candidate:
            return
        path = Path(candidate)
        key = str(path).lower() if platform.system() == "Windows" else str(path)
        if key in seen:
            return
        seen.add(key)
        candidates.append(path)

    java_home = os.environ.get("JAVA_HOME")
    if java_home:
        add(Path(java_home) / "bin" / executable)
    add(shutil.which("java"))

    if platform.system() == "Windows":
        install_roots = [
            os.environ.get("ProgramW6432"),
            os.environ.get("ProgramFiles"),
            os.environ.get("ProgramFiles(x86)"),
        ]
        patterns = (
            "Eclipse Adoptium/*/bin/java.exe",
            "Java/*/bin/java.exe",
            "Microsoft/jdk-*/bin/java.exe",
            "Microsoft/*/bin/java.exe",
            "BellSoft/*/bin/java.exe",
            "Amazon Corretto/*/bin/java.exe",
        )
        for root in install_roots:
            if not root:
                continue
            base = Path(root)
            for pattern in patterns:
                for path in base.glob(pattern):
                    add(path)

    return candidates


def java_version_for_path(java_bin: Path) -> tuple[Optional[int], str]:
    try:
        result = subprocess.run(
            [str(java_bin), "-version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=8,
        )
    except Exception as exc:
        return None, compact_error_message(exc)
    output = "\n".join(part for part in (result.stderr, result.stdout) if part)
    if result.returncode != 0:
        return None, output.strip()
    return parse_java_major_version(output), output


def activate_java_runtime(java_bin: Path) -> None:
    bin_dir = str(java_bin.parent)
    current_path = os.environ.get("PATH", "")
    path_parts = current_path.split(os.pathsep) if current_path else []
    if not any(part.lower() == bin_dir.lower() for part in path_parts):
        os.environ["PATH"] = bin_dir + (os.pathsep + current_path if current_path else "")
    os.environ["JAVA_HOME"] = str(java_bin.parent.parent)


def java_runtime_error() -> str:
    candidates = java_candidate_paths()
    if not candidates:
        return f"Falta Java {MIN_LANGUAGETOOL_JAVA_MAJOR} o superior. {java_install_hint()}"

    checked: list[tuple[Path, Optional[int], str]] = []
    for candidate in candidates:
        major, output = java_version_for_path(candidate)
        checked.append((candidate, major, output))
        if major is not None and major >= MIN_LANGUAGETOOL_JAVA_MAJOR:
            activate_java_runtime(candidate)
            return ""

    detected = next((item for item in checked if item[1] is not None), None)
    if detected is not None:
        java_bin, major, _ = detected
        return (
            f"Java {major} detectado en {java_bin}. "
            f"LanguageTool {LANGUAGETOOL_VERSION} requiere Java {MIN_LANGUAGETOOL_JAVA_MAJOR} o superior. "
            f"{java_install_hint()}"
        )
    detail = checked[0][2].strip().splitlines()[0] if checked and checked[0][2].strip() else str(candidates[0])
    return f"No pude detectar la version de Java. LanguageTool requiere Java {MIN_LANGUAGETOOL_JAVA_MAJOR} o superior. {detail[:160]}"


def proofread_status_snapshot() -> dict[str, Any]:
    with PROOFREAD_STATUS_LOCK:
        return dict(PROOFREAD_STATUS)


def set_proofread_status(status: str, message: str = "", missing: Optional[list[str]] = None) -> None:
    with PROOFREAD_STATUS_LOCK:
        PROOFREAD_STATUS.update(
            {
                "status": status,
                "available": status == "ready",
                "message": message,
                "missing": missing or [],
                "updated_at": now_ms(),
            }
        )


def language_tool_dir_has_jar(directory: Path) -> bool:
    return any(path.is_file() for pattern in LANGUAGETOOL_JAR_PATTERNS for path in directory.glob(pattern))


def latest_language_tool_dir(base_dir: Path) -> Optional[Path]:
    if not base_dir.is_dir():
        return None
    candidates = [
        path
        for path in base_dir.glob("LanguageTool*")
        if path.is_dir() and language_tool_dir_has_jar(path)
    ]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.name)


def configure_language_tool_paths() -> None:
    LANGUAGETOOL_DIR.mkdir(parents=True, exist_ok=True)
    if os.environ.get("LTP_JAR_DIR_PATH"):
        return
    for base_dir in (LANGUAGETOOL_DIR, Path.home() / ".cache" / "language_tool_python"):
        existing = latest_language_tool_dir(base_dir)
        if existing is not None:
            os.environ["LTP_JAR_DIR_PATH"] = str(existing)
            os.environ.setdefault("LTP_PATH", str(base_dir))
            return
    os.environ.setdefault("LTP_PATH", str(LANGUAGETOOL_DIR))


def proofread_package_error() -> str:
    try:
        import language_tool_python  # type: ignore  # noqa: F401
    except ModuleNotFoundError:
        setup_script = "setup_windows.bat" if platform.system() == "Windows" else "setup_mac.sh"
        return f"Falta language-tool-python. Ejecuta {setup_script} o pip install -r requirements.txt."
    return ""


def proofread_start_error() -> str:
    package_error = proofread_package_error()
    if package_error:
        return package_error
    return java_runtime_error()


def start_proofread_background() -> None:
    global PROOFREAD_INIT_STARTED
    startup_error = proofread_start_error()
    if startup_error:
        set_proofread_status("unavailable", startup_error, [startup_error])
        return
    with PROOFREAD_STATUS_LOCK:
        if PROOFREAD_INIT_STARTED or PROOFREAD_STATUS.get("status") in {"preparing", "ready"}:
            return
        PROOFREAD_INIT_STARTED = True
        PROOFREAD_STATUS.update(
            {
                "status": "preparing",
                "available": False,
                "message": "Iniciando corrector local...",
                "missing": [],
                "updated_at": now_ms(),
            }
        )

    def worker() -> None:
        global PROOFREAD_INIT_STARTED
        try:
            get_proofread_tool("es")
            set_proofread_status("ready", "Corrector local listo.")
        except Exception as exc:
            set_proofread_status("unavailable", compact_error_message(exc), [compact_error_message(exc)])
            with PROOFREAD_STATUS_LOCK:
                PROOFREAD_INIT_STARTED = False

    threading.Thread(target=worker, daemon=True).start()


def stop_proofread_tool() -> None:
    global PROOFREAD_TOOL, PROOFREAD_INIT_STARTED
    with PROOFREAD_LOCK:
        tool = PROOFREAD_TOOL
        PROOFREAD_TOOL = None
    if tool is not None:
        close = getattr(tool, "close", None)
        if callable(close):
            try:
                close()
            except Exception:
                pass
    with PROOFREAD_STATUS_LOCK:
        PROOFREAD_INIT_STARTED = False
        PROOFREAD_STATUS.update(
            {
                "status": "idle",
                "available": False,
                "message": "Corrector local apagado.",
                "missing": [],
                "updated_at": now_ms(),
            }
        )


def proofread_cache_key(language: str, text: str) -> str:
    normalized_language = "es" if not language or language.lower().startswith("es") else language
    digest = hashlib.sha256(f"{normalized_language}\0{text}".encode("utf-8")).hexdigest()
    return digest


def cached_proofread_result(language: str, text: str) -> Optional[dict[str, Any]]:
    key = proofread_cache_key(language, text)
    with PROOFREAD_CACHE_LOCK:
        result = PROOFREAD_CACHE.get(key)
        return dict(result) if result is not None else None


def remember_proofread_result(language: str, text: str, result: dict[str, Any]) -> None:
    key = proofread_cache_key(language, text)
    with PROOFREAD_CACHE_LOCK:
        if key not in PROOFREAD_CACHE:
            PROOFREAD_CACHE_ORDER.append(key)
        PROOFREAD_CACHE[key] = dict(result)
        while len(PROOFREAD_CACHE_ORDER) > PROOFREAD_CACHE_LIMIT:
            old_key = PROOFREAD_CACHE_ORDER.pop(0)
            PROOFREAD_CACHE.pop(old_key, None)


def get_proofread_tool(language: str = "es") -> Any:
    global PROOFREAD_TOOL
    normalized_language = "es" if not language or language.lower().startswith("es") else language
    with PROOFREAD_LOCK:
        if PROOFREAD_TOOL is not None:
            return PROOFREAD_TOOL
        startup_error = proofread_start_error()
        if startup_error:
            raise RuntimeError(startup_error)
        configure_language_tool_paths()
        import language_tool_python  # type: ignore
        try:
            PROOFREAD_TOOL = language_tool_python.LanguageTool(
                normalized_language,
                language_tool_download_version=LANGUAGETOOL_VERSION,
            )
        except Exception as exc:
            raise RuntimeError(f"No pude iniciar LanguageTool local: {compact_error_message(exc)}") from exc
        return PROOFREAD_TOOL


def proofread_text(text: str, language: str = "es") -> dict[str, Any]:
    original = text or ""
    if not original.strip():
        return {"ok": True, "language": "es", "matches": [], "truncated": False, "cached": False}
    if len(original) > 8000:
        raise HTTPException(status_code=400, detail="Texto demasiado largo para revisar en una sola pasada.")
    cached = cached_proofread_result(language, original)
    if cached is not None:
        cached["cached"] = True
        return cached
    tool = get_proofread_tool(language)
    matches = tool.check(original)
    serialized = [serialize_proofread_match(match, original) for match in matches[:40]]
    result = {
        "ok": True,
        "language": "es",
        "matches": serialized,
        "truncated": len(matches) > len(serialized),
        "cached": False,
    }
    remember_proofread_result(language, original, result)
    return result


def serialize_proofread_match(match: Any, text: str) -> dict[str, Any]:
    offset = max(0, int(getattr(match, "offset", 0) or 0))
    length = max(0, int(getattr(match, "errorLength", 0) or 0))
    issue_text = text[offset : offset + length]
    replacements = [str(item) for item in (getattr(match, "replacements", None) or [])[:6]]
    category = getattr(getattr(match, "category", None), "name", None) or getattr(match, "category", None)
    return {
        "offset": offset,
        "length": length,
        "text": issue_text,
        "message": str(getattr(match, "message", "") or "Revisar texto"),
        "short_message": str(getattr(match, "shortMessage", "") or ""),
        "rule_id": str(getattr(match, "ruleId", "") or ""),
        "category": str(category or ""),
        "issue_type": str(getattr(match, "ruleIssueType", "") or ""),
        "replacements": replacements,
    }


def is_fatal_whisper_error(text: str) -> bool:
    lowered = text.lower()
    fatal_markers = [
        "failed to initialize whisper context",
        "failed to load model",
        "unknown tensor",
        "modelo whisper invalido",
        "no encontre whisper.cpp",
        "no hay modelos whisper",
    ]
    return any(marker in lowered for marker in fatal_markers)


def process_project(
    project_id: str,
    model_name: str,
    diarize: bool,
    num_speakers: Optional[int],
    profile: str,
    reset_manifest: bool = False,
) -> None:
    try:
        profile_name = normalize_processing_profile(profile)
        append_project_log(project_id, f"Iniciando procesamiento con modelo {model_name}, perfil {profile_name}")
        project = load_project(project_id)
        existing_labels = project.get("speaker_labels") or {}
        tools = get_tools()
        project["status"] = "processing"
        project["updated_at"] = now_ms()
        project["warnings"] = []
        project["error"] = None
        project["run_config"] = {
            "model": model_name,
            "diarize": diarize,
            "num_speakers": num_speakers,
            "profile": profile_name,
        }
        save_project(project)
        update_processing_manifest(project_id, stage="convert")
        update_job(project_id, status="processing", step="Convirtiendo audio", progress=10, stage="convert", can_resume=True)

        existing_wav = project_media_path(project, "audio_path", "audio")
        if existing_wav and existing_wav.is_file():
            wav_path = existing_wav
        else:
            wav_path = convert_audio(project, tools)
        project["audio_path"] = str(wav_path)
        save_project(project)
        ensure_not_stopped(project_id)

        update_job(project_id, step="Transcribiendo con Whisper", progress=25, stage="whisper", can_resume=True)
        segments = run_whisper(project, wav_path, tools, model_name, profile_name, reset_manifest=reset_manifest)
        project["segments"] = segments
        project["speaker_labels"] = default_speaker_labels(segments, existing_labels)
        project["model"] = model_name
        project["status"] = "transcribed"
        project["updated_at"] = now_ms()
        save_project(project)

        turns = []
        if diarize:
            ready = diarization_ready()
            if ready["ready"]:
                update_processing_manifest(project_id, stage="diarization")
                update_job(project_id, step="Separando hablantes", progress=70, stage="diarization", can_resume=True)
                ensure_not_stopped(project_id)
                try:
                    turns = run_diarization(
                        wav_path,
                        project_dir(project_id) / "diarization_turns.json",
                        num_speakers=num_speakers,
                    )
                except Exception as exc:
                    project["warnings"].append(f"Diarizacion omitida: {compact_error_message(exc)}")
            else:
                project["warnings"].append("Diarizacion no disponible. Revisa setup o requirements_diarization.txt.")

        ensure_not_stopped(project_id)
        update_job(project_id, step="Preparando editor", progress=90, stage="editor", can_resume=False)
        if turns:
            segments, speaker_labels, _, diarization_warning = relabel_segments_with_diagnostics(
                project_id,
                segments,
                turns,
                existing_labels,
            )
            if diarization_warning:
                project["warnings"] = clear_diarization_warnings(project.get("warnings") or [])
                project["warnings"].append(diarization_warning)
        else:
            speaker_labels = default_speaker_labels(segments, existing_labels)
        project["segments"] = segments
        project["diarization_turns"] = turns
        project["speaker_labels"] = speaker_labels
        project["model"] = model_name
        project["status"] = "done"
        bump_content_revision(project)
        project["updated_at"] = now_ms()
        save_project(project)
        update_processing_manifest(project_id, stage="done")
        update_job(project_id, status="done", step="Listo", progress=100, stage="done", can_resume=False)
    except JobStopped as exc:
        try:
            project = load_project(project_id)
            project["status"] = exc.status
            project["error"] = None
            project["updated_at"] = now_ms()
            save_project(project)
            update_processing_manifest(project_id, stage=exc.status)
        except Exception:
            pass
        step = "Pausado. Puedes reanudar cuando quieras." if exc.status == "paused" else "Cancelado."
        update_job(project_id, status=exc.status, step=step, progress=100 if exc.status == "cancelled" else 0, can_resume=True)
    except Exception as exc:
        try:
            project = load_project(project_id)
            project["status"] = "error"
            project["error"] = str(exc)
            project["updated_at"] = now_ms()
            save_project(project)
            update_processing_manifest(project_id, stage="error", error=str(exc))
        except Exception:
            pass
        update_job(project_id, status="error", step="Error", progress=100, error=str(exc), can_resume=True)


def format_timestamp(seconds: float, sep: str = ",") -> str:
    seconds = max(0.0, float(seconds))
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    whole = int(seconds % 60)
    millis = int(round((seconds - int(seconds)) * 1000))
    if millis == 1000:
        whole += 1
        millis = 0
    return f"{hours:02d}:{minutes:02d}:{whole:02d}{sep}{millis:03d}"


def format_clock(seconds: float) -> str:
    seconds = max(0.0, float(seconds))
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    whole = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{whole:02d}"


def speaker_name(project: dict[str, Any], speaker: str) -> str:
    labels = project.get("speaker_labels") or {}
    return labels.get(speaker, speaker)


def start_processing_thread(
    project_id: str,
    model_name: str,
    diarize: bool,
    num_speakers: Optional[int],
    profile: str,
    *,
    reset_manifest: bool = False,
) -> None:
    project = load_project(project_id)
    if project.get("status") in {"queued", "processing"}:
        raise HTTPException(status_code=400, detail="Ese proyecto ya se esta procesando.")
    clear_stop_request(project_id)
    profile_name = normalize_processing_profile(profile)
    project["status"] = "queued"
    project["error"] = None
    project["updated_at"] = now_ms()
    project["run_config"] = {
        "model": model_name,
        "diarize": diarize,
        "num_speakers": num_speakers,
        "profile": profile_name,
    }
    save_project(project)
    update_job(project_id, status="queued", step="En cola", progress=0, started_at=now_ms(), can_resume=True)
    thread = threading.Thread(
        target=process_project,
        args=(project_id, model_name, diarize, num_speakers, profile_name, reset_manifest),
        daemon=True,
    )
    thread.start()


def process_diarization_only(project_id: str, num_speakers: Optional[int]) -> None:
    try:
        append_project_log(project_id, f"Iniciando diarizacion aislada. speakers={num_speakers or 'auto'}")
        clear_stop_request(project_id)
        project = load_project(project_id)
        existing_labels = project.get("speaker_labels") or {}
        if not project.get("segments"):
            raise RuntimeError("No hay segmentos transcritos para diarizar.")
        audio_path = project_media_path(project, "audio_path", "audio")
        if not audio_path or not audio_path.is_file():
            raise RuntimeError("No hay WAV convertido para diarizar.")
        if not diarization_ready()["ready"]:
            raise RuntimeError("Diarizacion no disponible. Revisa setup o requirements_diarization.txt.")

        project["status"] = "processing"
        project["error"] = None
        project["warnings"] = clear_diarization_warnings(project.get("warnings") or [])
        project["updated_at"] = now_ms()
        save_project(project)
        update_processing_manifest(project_id, stage="diarization")
        update_job(
            project_id,
            status="processing",
            step="Separando hablantes",
            progress=70,
            started_at=now_ms(),
            stage="diarization",
            can_resume=True,
        )

        turns = run_diarization(
            audio_path,
            project_dir(project_id) / "diarization_turns.json",
            num_speakers=num_speakers,
        )
        segments, speaker_labels, _, diarization_warning = relabel_segments_with_diagnostics(
            project_id,
            project.get("segments") or [],
            turns,
            existing_labels,
        )
        project = load_project(project_id)
        project["segments"] = segments
        project["diarization_turns"] = turns
        project["speaker_labels"] = speaker_labels
        project["status"] = "done"
        project["error"] = None
        bump_content_revision(project)
        project["updated_at"] = now_ms()
        warnings = project.get("warnings") or []
        project["warnings"] = clear_diarization_warnings(warnings)
        if diarization_warning:
            project["warnings"].append(diarization_warning)
        save_project(project)
        update_processing_manifest(project_id, stage="done")
        update_job(project_id, status="done", step="Listo", progress=100, stage="done", can_resume=False)
    except Exception as exc:
        try:
            project = load_project(project_id)
            project["status"] = "done" if project.get("segments") else "error"
            project["error"] = None if project.get("segments") else str(exc)
            warnings = clear_diarization_warnings(project.get("warnings") or [])
            warnings.append(f"Diarizacion omitida: {compact_error_message(exc)}")
            project["warnings"] = warnings
            project["updated_at"] = now_ms()
            save_project(project)
        except Exception:
            pass
        update_job(project_id, status="done", step="Diarizacion omitida", progress=100)


def start_diarization_thread(project_id: str, num_speakers: Optional[int]) -> None:
    project = load_project(project_id)
    if project.get("status") in {"queued", "processing"}:
        raise HTTPException(status_code=400, detail="Ese proyecto ya se esta procesando.")
    project["status"] = "queued"
    project["error"] = None
    project["updated_at"] = now_ms()
    save_project(project)
    update_job(project_id, status="queued", step="Preparando diarizacion", progress=0, started_at=now_ms(), stage="diarization", can_resume=True)
    thread = threading.Thread(
        target=process_diarization_only,
        args=(project_id, num_speakers),
        daemon=True,
    )
    thread.start()


def export_txt(project: dict[str, Any]) -> str:
    lines = []
    last_speaker = None
    buffer: list[str] = []
    for segment in project.get("segments", []):
        current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
        if current != last_speaker and buffer:
            lines.append(f"{last_speaker}:")
            lines.append(" ".join(buffer).strip())
            lines.append("")
            buffer = []
        last_speaker = current
        buffer.append((segment.get("text") or "").strip())
    if buffer:
        lines.append(f"{last_speaker}:")
        lines.append(" ".join(buffer).strip())
    return "\n".join(lines).strip() + "\n"


def export_srt(project: dict[str, Any], vtt: bool = False) -> str:
    lines = ["WEBVTT", ""] if vtt else []
    for idx, segment in enumerate(project.get("segments", []), start=1):
        if not vtt:
            lines.append(str(idx))
        start = format_timestamp(segment.get("start", 0), "." if vtt else ",")
        end = format_timestamp(segment.get("end", 0), "." if vtt else ",")
        lines.append(f"{start} --> {end}")
        lines.append(f"{speaker_name(project, segment.get('speaker', 'SPEAKER_00'))}: {segment.get('text', '').strip()}")
        lines.append("")
    return "\n".join(lines)


def export_docx(project: dict[str, Any], output_path: Path, include_timestamps: bool = False) -> None:
    document = Document()
    document.add_heading(project.get("name") or "Transcripcion", level=1)
    if include_timestamps:
        for segment in project.get("segments", []):
            current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
            start = format_clock(segment.get("start", 0))
            end = format_clock(segment.get("end", 0))
            paragraph = document.add_paragraph()
            time_run = paragraph.add_run(f"[{start}-{end}] ")
            time_run.italic = True
            speaker_run = paragraph.add_run(f"{current}: ")
            speaker_run.bold = True
            paragraph.add_run((segment.get("text") or "").strip())
        document.save(output_path)
        return

    last_speaker = None
    paragraph = None
    for segment in project.get("segments", []):
        current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
        if current != last_speaker:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{current}: ")
            run.bold = True
            last_speaker = current
        if paragraph is None:
            paragraph = document.add_paragraph()
        paragraph.add_run((segment.get("text") or "").strip() + " ")
    document.save(output_path)


def truthy_param(value: Any) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "si", "sí", "on"}


def safe_zip_member(name: str) -> str:
    raw = str(name or "")
    if not raw or "\x00" in raw:
        raise HTTPException(status_code=400, detail="Paquete invalido: ruta vacia.")
    if raw.startswith(("/", "\\")) or re.match(r"^[A-Za-z]:[\\/]", raw):
        raise HTTPException(status_code=400, detail="Paquete invalido: ruta no permitida.")
    normalized = raw.replace("\\", "/")
    parts = normalized.split("/")
    if any(part in {"", ".", ".."} for part in parts):
        raise HTTPException(status_code=400, detail="Paquete invalido: ruta no permitida.")
    return "/".join(parts)


def package_member_allowed(member: str) -> bool:
    if member in {"project.json", "manifest.json"}:
        return True
    parts = member.split("/")
    if len(parts) != 2:
        return False
    if parts[0] == "audio":
        suffix = Path(parts[1]).suffix.lower() or ".audio"
        return suffix in ALLOWED_IMPORT_AUDIO_SUFFIXES
    if parts[0] == "diagnostics":
        return parts[1] in ALLOWED_PACKAGE_DIAGNOSTICS
    return False


def is_zip_symlink(info: zipfile.ZipInfo) -> bool:
    return ((info.external_attr >> 16) & 0o170000) == 0o120000


def validated_package_entries(package: zipfile.ZipFile) -> dict[str, zipfile.ZipInfo]:
    entries: dict[str, zipfile.ZipInfo] = {}
    total_size = 0
    for info in package.infolist():
        if info.is_dir():
            continue
        if is_zip_symlink(info):
            raise HTTPException(status_code=400, detail="Paquete invalido: enlaces simbolicos no permitidos.")
        member = safe_zip_member(info.filename)
        if not package_member_allowed(member):
            raise HTTPException(status_code=400, detail=f"Paquete invalido: archivo no permitido ({member}).")
        if member in entries:
            raise HTTPException(status_code=400, detail=f"Paquete invalido: archivo duplicado ({member}).")
        total_size += int(info.file_size or 0)
        if total_size > MAX_IMPORT_PACKAGE_BYTES:
            raise HTTPException(status_code=400, detail="Paquete invalido: contenido demasiado grande.")
        entries[member] = info
        if len(entries) > MAX_PACKAGE_MEMBERS:
            raise HTTPException(status_code=400, detail="Paquete invalido: demasiados archivos.")
    if "project.json" not in entries:
        raise HTTPException(status_code=400, detail="Paquete invalido: falta project.json.")
    return entries


def portable_project(project: dict[str, Any]) -> dict[str, Any]:
    portable = json.loads(json.dumps(project, ensure_ascii=False))
    portable["portable_format"] = "transcriptor-local-project"
    portable["portable_version"] = 1
    portable["original_id"] = project.get("id")
    portable["exported_at"] = now_ms()
    portable["source_path"] = ""
    portable["audio_path"] = ""
    portable["status"] = "done" if portable.get("segments") else portable.get("status", "done")
    return portable


def project_audio_for_package(project: dict[str, Any]) -> tuple[Optional[Path], str]:
    project_id = project.get("id")
    if not project_id:
        return None, ""
    source = project_media_path(project, "source_path", "source")
    if source and source.exists() and source.is_file():
        suffix = source.suffix or Path(project.get("source_name") or "").suffix or ".audio"
        return source, f"audio/source{suffix}"
    audio = project_media_path(project, "audio_path", "audio")
    if audio and audio.exists() and audio.is_file():
        return audio, f"audio/audio{audio.suffix or '.wav'}"
    return None, ""


def add_file_if_exists(package: zipfile.ZipFile, path: Path, arcname: str) -> bool:
    if path.exists() and path.is_file():
        package.write(path, arcname)
        return True
    return False


def export_package(project: dict[str, Any], output_path: Path, include_audio: bool = True) -> None:
    pdir = project_dir(project["id"])
    portable = portable_project(project)
    manifest: dict[str, Any] = {
        "format": "transcriptor-local-package",
        "version": 1,
        "created_at": now_ms(),
        "app": "Transcriptor Mi Cami",
        "project": {
            "id": project.get("id"),
            "name": project.get("name"),
            "segments": len(project.get("segments") or []),
            "speakers": len(project.get("speaker_labels") or {}),
            "has_diarization": bool(project.get("diarization_turns")),
        },
        "audio": None,
    }

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as package:
        if include_audio:
            audio_path, audio_arcname = project_audio_for_package(project)
            if audio_path and audio_arcname:
                package.write(audio_path, audio_arcname)
                manifest["audio"] = {"path": audio_arcname, "source_name": project.get("source_name") or audio_path.name}
                if audio_arcname.startswith("audio/source"):
                    portable["source_path"] = audio_arcname
                    portable["audio_path"] = ""
                else:
                    portable["source_path"] = ""
                    portable["audio_path"] = audio_arcname

        package.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        package.writestr("project.json", json.dumps(portable, ensure_ascii=False, indent=2))

        for filename in (
            "whisper_quality.json",
            "diarization_quality.json",
            "diarization_turns.json",
            "process.log",
        ):
            add_file_if_exists(package, pdir / filename, f"diagnostics/{filename}")


def read_package_json(package: zipfile.ZipFile, entries: dict[str, zipfile.ZipInfo], member: str) -> dict[str, Any]:
    try:
        info = entries[member]
    except KeyError as exc:
        raise HTTPException(status_code=400, detail=f"Paquete invalido: falta {member}.") from exc
    if info.file_size > MAX_PACKAGE_JSON_BYTES:
        raise HTTPException(status_code=400, detail=f"Paquete invalido: {member} es demasiado grande.")
    try:
        with package.open(info) as file:
            payload = json.loads(file.read().decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=400, detail=f"Paquete invalido: {member} no es JSON valido.") from exc
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail=f"Paquete invalido: {member} debe ser objeto JSON.")
    return payload


def copy_package_member(package: zipfile.ZipFile, info: zipfile.ZipInfo, destination: Path, max_bytes: int, description: str) -> None:
    if info.file_size > max_bytes:
        raise HTTPException(status_code=400, detail=f"{description} del paquete es demasiado grande.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    with package.open(info) as source, destination.open("wb") as target:
        shutil.copyfileobj(source, target)


async def write_upload_to_path(upload: UploadFile, destination: Path, max_bytes: int, description: str) -> int:
    destination.parent.mkdir(parents=True, exist_ok=True)
    uploaded_bytes = 0
    try:
        with destination.open("wb") as out:
            while True:
                chunk = await upload.read(1024 * 1024)
                if not chunk:
                    break
                uploaded_bytes += len(chunk)
                if uploaded_bytes > max_bytes:
                    raise HTTPException(status_code=400, detail=f"{description} es demasiado grande.")
                out.write(chunk)
        return uploaded_bytes
    except Exception:
        destination.unlink(missing_ok=True)
        raise


def inspect_project_package(package_path: Path) -> dict[str, Any]:
    try:
        package = zipfile.ZipFile(package_path)
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400, detail="Paquete invalido: no es un ZIP valido.") from exc
    with package:
        entries = validated_package_entries(package)
        project = read_package_json(package, entries, "project.json")
        manifest = read_package_json(package, entries, "manifest.json") if "manifest.json" in entries else {}
        preview_project = json.loads(json.dumps(project, ensure_ascii=False))
        normalize_imported_segments(preview_project)
        segments = preview_project.get("segments") if isinstance(preview_project.get("segments"), list) else []
        labels = preview_project.get("speaker_labels") if isinstance(preview_project.get("speaker_labels"), dict) else {}
        turns = preview_project.get("diarization_turns") if isinstance(preview_project.get("diarization_turns"), list) else []
        manifest_project = manifest.get("project") if isinstance(manifest.get("project"), dict) else {}
        manifest_audio = manifest.get("audio") if isinstance(manifest.get("audio"), dict) else {}
        audio_member = safe_zip_member(str(manifest_audio.get("path") or "")) if manifest_audio.get("path") else ""
        audio_entry = entries.get(audio_member) if audio_member else None
    original_id = project.get("original_id") or project.get("id")
    return {
        "original_id": str(original_id or ""),
        "name": clean_filename(str(project.get("name") or "Transcripcion importada")),
        "package_sha256": sha256_file(package_path),
        "manifest": manifest,
        "segments": len(segments),
        "speakers": len(labels) or len({str(segment.get("speaker") or "") for segment in segments if isinstance(segment, dict)}),
        "has_diarization": bool(turns or manifest_project.get("has_diarization")),
        "has_audio": bool(audio_entry),
        "audio_name": str(manifest_audio.get("source_name") or Path(audio_member).name) if audio_member else "",
        "audio_bytes": int(audio_entry.file_size) if audio_entry else 0,
        "created_at": project.get("created_at") or manifest.get("created_at"),
        "updated_at": project.get("updated_at") or project.get("exported_at") or manifest.get("created_at"),
    }


def find_duplicate_import(package_info: dict[str, Any]) -> Optional[dict[str, Any]]:
    original_id = package_info.get("original_id")
    package_sha256 = package_info.get("package_sha256")
    for project in list_projects():
        imported_from = project.get("imported_from") if isinstance(project.get("imported_from"), dict) else {}
        if package_sha256 and imported_from.get("package_sha256") == package_sha256:
            return project
        if original_id and (project.get("id") == original_id or imported_from.get("original_id") == original_id):
            return project
    return None


def unique_project_name(base_name: str) -> str:
    existing = {str(project.get("name") or "") for project in list_projects()}
    if base_name not in existing:
        return base_name
    first_copy = f"{base_name} (copia)"
    if first_copy not in existing:
        return first_copy
    index = 2
    while True:
        candidate = f"{base_name} (copia {index})"
        if candidate not in existing:
            return candidate
        index += 1


def normalize_imported_segments(project: dict[str, Any]) -> None:
    segments = project.get("segments")
    if not isinstance(segments, list):
        raise HTTPException(status_code=400, detail="Paquete invalido: no contiene segmentos.")
    if len(segments) > MAX_IMPORTED_SEGMENTS:
        raise HTTPException(status_code=400, detail="Paquete invalido: demasiados segmentos.")
    for index, segment in enumerate(segments):
        if not isinstance(segment, dict):
            raise HTTPException(status_code=400, detail="Paquete invalido: segmento invalido.")
        segment.setdefault("id", f"seg-{index:05d}")
        try:
            start = float(segment.get("start") or 0)
            end = float(segment.get("end") or start)
        except (TypeError, ValueError) as exc:
            raise HTTPException(status_code=400, detail="Paquete invalido: timestamps invalidos.") from exc
        if not math.isfinite(start) or not math.isfinite(end) or start < 0 or end < 0:
            raise HTTPException(status_code=400, detail="Paquete invalido: timestamps invalidos.")
        if end < start:
            raise HTTPException(status_code=400, detail="Paquete invalido: segmento con fin antes del inicio.")
        segment["id"] = str(segment.get("id") or f"seg-{index:05d}")[:80]
        segment["start"] = start
        segment["end"] = end
        segment["speaker"] = str(segment.get("speaker") or "SPEAKER_00")[:80]
        segment["text"] = str(segment.get("text") or "")[:MAX_IMPORTED_TEXT_CHARS]
    labels = project.get("speaker_labels")
    if not isinstance(labels, dict):
        project["speaker_labels"] = default_speaker_labels(segments, {})
    else:
        project["speaker_labels"] = {str(key)[:80]: str(value)[:120] for key, value in list(labels.items())[:1000]}
    turns = project.get("diarization_turns")
    if turns is not None and not isinstance(turns, list):
        project["diarization_turns"] = []
    elif isinstance(turns, list):
        normalized_turns = []
        for turn in turns[: MAX_IMPORTED_SEGMENTS * 3]:
            if not isinstance(turn, dict):
                continue
            try:
                start = float(turn.get("start") or 0)
                end = float(turn.get("end") or start)
            except (TypeError, ValueError):
                continue
            if not math.isfinite(start) or not math.isfinite(end) or start < 0 or end < start:
                continue
            normalized_turns.append(
                {
                    "start": start,
                    "end": end,
                    "speaker": str(turn.get("speaker") or "SPEAKER_00")[:80],
                }
            )
        project["diarization_turns"] = normalized_turns


def import_project_package(package_path: Path, package_info: Optional[dict[str, Any]] = None, copy_name: bool = False) -> dict[str, Any]:
    try:
        package = zipfile.ZipFile(package_path)
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400, detail="Paquete invalido: no es un ZIP valido.") from exc

    with package:
        entries = validated_package_entries(package)
        project = read_package_json(package, entries, "project.json")
        manifest = read_package_json(package, entries, "manifest.json") if "manifest.json" in entries else {}
        normalize_imported_segments(project)
        manifest_audio = manifest.get("audio") if isinstance(manifest.get("audio"), dict) else {}
        audio_info = manifest_audio
        audio_member = safe_zip_member(str(audio_info.get("path") or "")) if audio_info.get("path") else ""

        project_id = uuid.uuid4().hex[:12]
        pdir = project_dir(project_id)
        pdir.mkdir(parents=True, exist_ok=True)
        try:
            original_id = (package_info or {}).get("original_id") or project.get("original_id") or project.get("id")
            base_name = clean_filename(str(project.get("name") or (package_info or {}).get("name") or "Transcripcion importada"))
            project["id"] = project_id
            project["name"] = unique_project_name(base_name) if copy_name else base_name
            project["source_name"] = str(project.get("source_name") or manifest_audio.get("source_name") or "")
            project["status"] = "done"
            project["content_revision"] = project_content_revision(project)
            project["created_at"] = now_ms()
            project["updated_at"] = now_ms()
            project["imported_at"] = now_ms()
            project["imported_from"] = {
                "original_id": original_id,
                "package_format": manifest.get("format") or "transcriptor-local-package",
                "package_sha256": (package_info or {}).get("package_sha256") or sha256_file(package_path),
            }
            project["error"] = None
            project["warnings"] = project.get("warnings") if isinstance(project.get("warnings"), list) else []
            project["source_path"] = ""
            project["audio_path"] = ""

            if audio_member and audio_member in entries:
                suffix = Path(audio_member).suffix or ".audio"
                destination = pdir / f"source{suffix}"
                copy_package_member(package, entries[audio_member], destination, MAX_PACKAGE_AUDIO_BYTES, "El audio")
                project["source_path"] = str(destination)
                project["source_name"] = project.get("source_name") or Path(audio_member).name

            for filename in ("whisper_quality.json", "diarization_quality.json", "diarization_turns.json", "process.log"):
                member = f"diagnostics/{filename}"
                if member in entries:
                    copy_package_member(package, entries[member], pdir / filename, MAX_PACKAGE_DIAGNOSTIC_BYTES, "El diagnostico")

            save_project(project)
            append_project_log(project_id, f"Proyecto importado desde paquete portable. Original: {original_id or 'desconocido'}.")
            return project
        except Exception:
            shutil.rmtree(pdir, ignore_errors=True)
            raise


@app.get("/")
def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/api/status")
def api_status() -> dict[str, Any]:
    tools = get_tools()
    return {
        "tools": {
            "ffmpeg": bool(tools.ffmpeg),
            "ffmpeg_path": tools.ffmpeg,
            "whisper": bool(tools.whisper),
            "whisper_path": tools.whisper,
        },
        "models": available_models(),
        "profiles": [
            {"name": key, "label": value["label"], "default": key == default_processing_profile()}
            for key, value in PROCESSING_PROFILES.items()
        ],
        "default_profile": default_processing_profile(),
        "diarization": diarization_ready(),
        "data_dir": str(DATA_DIR),
    }


@app.get("/api/projects")
def api_projects() -> list[dict[str, Any]]:
    return list_projects()


@app.post("/api/projects")
async def api_create_project(
    file: UploadFile = File(...),
    model: str = Form("auto"),
    diarize: str = Form("true"),
    speakers: str = Form("auto"),
    profile: str = Form("auto"),
) -> dict[str, Any]:
    tools = get_tools()
    if not tools.ffmpeg or not tools.whisper:
        raise HTTPException(status_code=400, detail="Faltan herramientas. Ejecuta setup.")
    models = available_models()
    if not models:
        raise HTTPException(status_code=400, detail="No hay modelos Whisper. Ejecuta setup.")
    profile_name = normalize_processing_profile(profile)
    model_name = select_model_for_request(models, model, profile_name)
    if not any(item["name"] == model_name for item in models):
        raise HTTPException(status_code=400, detail="Modelo no disponible.")

    project_id = uuid.uuid4().hex[:12]
    pdir = project_dir(project_id)
    pdir.mkdir(parents=True, exist_ok=True)
    suffix = Path(file.filename or "audio").suffix
    source_path = pdir / f"source{suffix}"
    try:
        await write_upload_to_path(file, source_path, MAX_UPLOAD_AUDIO_BYTES, "El audio")
        num_speakers = parse_speaker_count(speakers)

        project = {
            "id": project_id,
            "name": clean_filename(file.filename or "audio"),
            "source_name": file.filename,
            "source_path": str(source_path),
            "status": "new",
            "content_revision": 0,
            "created_at": now_ms(),
            "updated_at": now_ms(),
            "segments": [],
            "speaker_labels": {},
            "warnings": [],
            "error": None,
            "run_config": {
                "model": model_name,
                "diarize": diarize.lower() == "true",
                "num_speakers": num_speakers,
                "profile": profile_name,
            },
        }
        save_project(project)
        start_processing_thread(project_id, model_name, diarize.lower() == "true", num_speakers, profile_name)
        return {"id": project_id, "status": "queued"}
    except Exception:
        shutil.rmtree(pdir, ignore_errors=True)
        raise


@app.post("/api/import/package")
async def api_import_package(file: UploadFile = File(...), duplicate_mode: str = Form("ask")) -> dict[str, Any]:
    filename = file.filename or "transcripcion.transcriptor.zip"
    if not filename.lower().endswith((".zip", ".transcriptor.zip")):
        raise HTTPException(status_code=400, detail="Selecciona un paquete .transcriptor.zip o .zip.")
    duplicate_mode = duplicate_mode.strip().lower()
    if duplicate_mode not in {"ask", "copy", "open"}:
        raise HTTPException(status_code=400, detail="Modo de duplicado no soportado.")
    cleanup_stale_export_temps()
    temp_path = temporary_export_path("import", ".zip")
    try:
        await write_upload_to_path(file, temp_path, MAX_IMPORT_PACKAGE_BYTES, "El paquete")
        package_info = inspect_project_package(temp_path)
        duplicate = find_duplicate_import(package_info)
        if duplicate and duplicate_mode in {"ask", "open"}:
            return {
                "duplicate": True,
                "existing": project_summary(duplicate),
                "package": {
                    "name": package_info.get("name"),
                    "original_id": package_info.get("original_id"),
                    "segments": package_info.get("segments"),
                    "speakers": package_info.get("speakers"),
                    "has_diarization": package_info.get("has_diarization"),
                    "has_audio": package_info.get("has_audio"),
                    "audio_name": package_info.get("audio_name"),
                    "audio_bytes": package_info.get("audio_bytes"),
                    "updated_at": package_info.get("updated_at"),
                },
            }
        project = import_project_package(temp_path, package_info=package_info, copy_name=duplicate_mode == "copy" and duplicate is not None)
        return {"id": project["id"], "status": project.get("status"), "project": project}
    finally:
        temp_path.unlink(missing_ok=True)


@app.post("/api/import/package/inspect")
async def api_import_package_inspect(file: UploadFile = File(...)) -> dict[str, Any]:
    filename = file.filename or "transcripcion.transcriptor.zip"
    if not filename.lower().endswith((".zip", ".transcriptor.zip")):
        raise HTTPException(status_code=400, detail="Selecciona un paquete .transcriptor.zip o .zip.")
    cleanup_stale_export_temps()
    temp_path = temporary_export_path("inspect", ".zip")
    try:
        await write_upload_to_path(file, temp_path, MAX_IMPORT_PACKAGE_BYTES, "El paquete")
        package_info = inspect_project_package(temp_path)
        duplicate = find_duplicate_import(package_info)
        return {
            "package": {
                "name": package_info.get("name"),
                "original_id": package_info.get("original_id"),
                "segments": package_info.get("segments"),
                "speakers": package_info.get("speakers"),
                "has_diarization": package_info.get("has_diarization"),
                "has_audio": package_info.get("has_audio"),
                "audio_name": package_info.get("audio_name"),
                "audio_bytes": package_info.get("audio_bytes"),
                "created_at": package_info.get("created_at"),
                "updated_at": package_info.get("updated_at"),
            },
            "duplicate": bool(duplicate),
            "existing": project_summary(duplicate) if duplicate else None,
        }
    finally:
        temp_path.unlink(missing_ok=True)


@app.get("/api/projects/{project_id}")
def api_project(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    return load_project(project_id)


@app.patch("/api/projects/{project_id}")
def api_update_project(project_id: str, payload: ProjectUpdate) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    if payload.name is not None:
        project["name"] = payload.name.strip() or project["name"]
    if payload.playback_position is not None:
        project["playback_position"] = max(0.0, float(payload.playback_position))
    project["updated_at"] = now_ms()
    save_project(project)
    return project


@app.get("/api/jobs/{project_id}")
def api_job(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    with JOBS_LOCK:
        job = public_job(dict(JOBS.get(project_id, {})))
    if not job:
        project = load_project(project_id)
        manifest = read_json_or_none(processing_manifest_path(project_id))
        job = {
            "status": project.get("status"),
            "step": project.get("status"),
            "progress": 100 if project.get("status") in {"done", "error"} else 0,
            "error": project.get("error"),
            "stage": manifest.get("stage") if isinstance(manifest, dict) else project.get("status"),
            "can_resume": project.get("status") in {"paused", "cancelled", "error"},
        }
    return job


@app.get("/api/projects/{project_id}/logs")
def api_project_logs(project_id: str) -> dict[str, str]:
    project_id = validate_project_id(project_id)
    load_project(project_id)
    return {"log": read_project_log(project_id)}


@app.post("/api/jobs/{project_id}/pause")
def api_pause_job(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    return request_job_stop(project_id, "paused")


@app.post("/api/jobs/{project_id}/cancel")
def api_cancel_job(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    return request_job_stop(project_id, "cancelled")


@app.post("/api/projects/{project_id}/resume")
def api_resume_project(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    if project.get("status") in {"queued", "processing"}:
        raise HTTPException(status_code=400, detail="Ese proyecto ya se esta procesando.")
    previous_status = project.get("status")
    config = project.get("run_config") or {}
    models = available_models()
    profile_name = normalize_processing_profile(config.get("profile"))
    model_name = select_model_for_request(models, config.get("model") or "auto", profile_name)
    diarize = bool(config.get("diarize", True))
    num_speakers = config.get("num_speakers")
    start_processing_thread(
        project_id,
        model_name,
        diarize,
        num_speakers,
        profile_name,
        reset_manifest=previous_status == "done",
    )
    return {"id": project_id, "status": "queued"}


@app.post("/api/projects/{project_id}/diarize")
def api_diarize_project(project_id: str, speakers: str = Form("auto")) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    num_speakers = parse_speaker_count(speakers)
    start_diarization_thread(project_id, num_speakers)
    return {"id": project_id, "status": "queued"}


@app.post("/api/projects/{project_id}/relabel-speakers")
def api_relabel_speakers(project_id: str, payload: Optional[RelabelRequest] = None) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    if project.get("status") in {"queued", "processing"}:
        raise HTTPException(status_code=400, detail="Ese proyecto se esta procesando.")
    if (payload and payload.mode) not in {None, "interview_2p"}:
        raise HTTPException(status_code=400, detail="Modo de reetiquetado no soportado.")
    segments = project.get("segments") or []
    turns = project.get("diarization_turns") or []
    if not segments:
        raise HTTPException(status_code=400, detail="No hay segmentos para reetiquetar.")
    if not turns:
        raise HTTPException(status_code=400, detail="No hay diarizacion guardada para reetiquetar.")

    existing_labels = project.get("speaker_labels") or {}
    segments, labels, _, diarization_warning = relabel_segments_with_diagnostics(
        project_id,
        segments,
        turns,
        existing_labels,
    )
    project["segments"] = segments
    project["speaker_labels"] = labels
    project["warnings"] = clear_diarization_warnings(project.get("warnings") or [])
    if diarization_warning:
        project["warnings"].append(diarization_warning)
    bump_content_revision(project)
    project["updated_at"] = now_ms()
    save_project(project)
    return project


@app.post("/api/projects/{project_id}/segments")
def api_save_segments(project_id: str, payload: SegmentUpdate) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    current_revision = project_content_revision(project)
    if payload.base_content_revision is not None and int(payload.base_content_revision) != current_revision:
        raise HTTPException(
            status_code=409,
            detail="Hay cambios guardados desde otra ventana. Revisa antes de sobrescribir.",
        )
    project["segments"] = payload.segments
    project["speaker_labels"] = payload.speaker_labels
    if payload.name is not None:
        project["name"] = payload.name.strip() or project.get("name") or "Transcripcion"
    bump_content_revision(project)
    project["updated_at"] = now_ms()
    save_project(project)
    return project


@app.get("/api/proofread/status")
def api_proofread_status(start: bool = False) -> dict[str, Any]:
    status = proofread_status_snapshot()
    if start and status.get("status") in {"idle", "unavailable"}:
        start_proofread_background()
        status = proofread_status_snapshot()
    return status


@app.post("/api/proofread/stop")
def api_proofread_stop() -> dict[str, Any]:
    stop_proofread_tool()
    return {"ok": True, **proofread_status_snapshot()}


@app.post("/api/proofread")
def api_proofread(payload: ProofreadRequest) -> dict[str, Any]:
    try:
        return proofread_text(payload.text, payload.language)
    except Exception as exc:
        raise HTTPException(status_code=503, detail=compact_error_message(exc)) from exc


@app.post("/api/proofread/batch")
def api_proofread_batch(payload: ProofreadBatchRequest) -> dict[str, Any]:
    if len(payload.items) > 24:
        raise HTTPException(status_code=400, detail="Demasiados segmentos para revisar en una sola pasada.")
    results = []
    try:
        for item in payload.items:
            result = proofread_text(item.text, payload.language)
            results.append(
                {
                    "id": item.id,
                    "matches": result.get("matches", []),
                    "truncated": bool(result.get("truncated")),
                    "cached": bool(result.get("cached")),
                }
            )
    except Exception as exc:
        raise HTTPException(status_code=503, detail=compact_error_message(exc)) from exc
    return {"ok": True, "language": "es", "results": results}


@app.get("/api/projects/{project_id}/audio")
def api_audio(project_id: str) -> FileResponse:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    audio_path = project_media_path(project, "audio_path", "audio")
    if audio_path and audio_path.is_file():
        return FileResponse(audio_path, media_type="audio/wav")
    source = project_media_path(project, "source_path", "source")
    if source and source.is_file():
        return FileResponse(source)
    raise HTTPException(status_code=404, detail="Audio no encontrado")


@app.get("/api/projects/{project_id}/export/{fmt}")
def api_export(project_id: str, fmt: str, audio: str = "true") -> Response:
    project_id = validate_project_id(project_id)
    project = load_project(project_id)
    safe_name = clean_filename(project.get("name") or "transcripcion")
    if fmt in {"package", "package-lite"}:
        include_audio = fmt == "package" and truthy_param(audio)
        suffix = "" if include_audio else "_sin_audio"
        output_path = safe_export_path(project_id, f"{suffix}.transcriptor.zip")
        write_export_atomically(output_path, lambda tmp: export_package(project, tmp, include_audio=include_audio))
        return FileResponse(
            output_path,
            media_type="application/zip",
            filename=f"{safe_name}{suffix}.transcriptor.zip",
        )
    if fmt == "txt":
        return PlainTextResponse(
            export_txt(project),
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.txt"'},
        )
    if fmt == "srt":
        return PlainTextResponse(
            export_srt(project),
            media_type="application/x-subrip",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.srt"'},
        )
    if fmt == "vtt":
        return PlainTextResponse(
            export_srt(project, vtt=True),
            media_type="text/vtt",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.vtt"'},
        )
    if fmt == "json":
        payload = json.dumps(project, ensure_ascii=False, indent=2)
        return JSONResponse(
            content=json.loads(payload),
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.json"'},
        )
    if fmt == "docx":
        output_path = safe_export_path(project_id, ".docx")
        write_export_atomically(output_path, lambda tmp: export_docx(project, tmp))
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{safe_name}.docx",
        )
    if fmt == "docx-ts":
        output_path = safe_export_path(project_id, "-timestamps.docx")
        write_export_atomically(output_path, lambda tmp: export_docx(project, tmp, include_timestamps=True))
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{safe_name}_timestamps.docx",
        )
    raise HTTPException(status_code=404, detail="Formato no soportado")


def delete_project_files(project_id: str) -> None:
    project_id = validate_project_id(project_id)
    pdir = project_dir(project_id)
    if not pdir.exists():
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    stop_project_for_delete(project_id)
    shutil.rmtree(pdir)
    for export_path in EXPORTS_DIR.glob(f"{project_id}*"):
        if export_path.is_file():
            export_path.unlink(missing_ok=True)
    for export_path in EXPORTS_DIR.glob(f"{EXPORT_TEMP_PREFIX}*{project_id}*"):
        if export_path.is_file():
            export_path.unlink(missing_ok=True)
    with JOBS_LOCK:
        JOBS.pop(project_id, None)
        RUNNING_PROCESSES.pop(project_id, None)


@app.delete("/api/projects/{project_id}")
def api_delete_project(project_id: str) -> dict[str, Any]:
    project_id = validate_project_id(project_id)
    delete_project_files(project_id)
    return {"ok": True}


def find_available_port(host: str, preferred_port: int, attempts: int = 25) -> int:
    for port in range(preferred_port, preferred_port + attempts):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                sock.bind((host, port))
            except OSError:
                continue
            return port
    raise RuntimeError(f"No encontre un puerto libre entre {preferred_port} y {preferred_port + attempts - 1}.")


def main() -> None:
    import uvicorn

    port = find_available_port(HOST, PORT)
    url = f"http://{HOST}:{port}"
    if port != PORT:
        print(f"Puerto {PORT} ocupado. Usando {port}.")
    print(f"Abriendo Transcriptor Mi Cami en {url}")
    if os.environ.get("TRANSCRIPTOR_NO_BROWSER") != "1":
        threading.Timer(1.0, lambda: webbrowser.open(url)).start()
    uvicorn.run("app.main:app", host=HOST, port=port, reload=False, access_log=False)


if __name__ == "__main__":
    main()
