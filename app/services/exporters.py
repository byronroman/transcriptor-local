from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def format_timestamp(seconds: float, sep: str = ",") -> str:
    seconds = max(0.0, float(seconds))
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    whole = int(seconds % 60)
    millis = int(round((seconds - int(seconds)) * 1000))
    if millis == 1000:
        whole += 1
        millis = 0
    return f"{hours:02d}:{minutes:02d}:{whole:02d}{sep}{millis:03d}"


def format_clock(seconds: float) -> str:
    seconds = max(0.0, float(seconds))
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    whole = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{whole:02d}"


def speaker_name(project: dict[str, Any], speaker: str) -> str:
    labels = project.get("speaker_labels") or {}
    return labels.get(speaker, speaker)


def export_txt(project: dict[str, Any]) -> str:
    lines = []
    last_speaker = None
    buffer: list[str] = []
    for segment in project.get("segments", []):
        current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
        if current != last_speaker and buffer:
            lines.append(f"{last_speaker}:")
            lines.append(" ".join(buffer).strip())
            lines.append("")
            buffer = []
        last_speaker = current
        buffer.append((segment.get("text") or "").strip())
    if buffer:
        lines.append(f"{last_speaker}:")
        lines.append(" ".join(buffer).strip())
    return "\n".join(lines).strip() + "\n"


def export_srt(project: dict[str, Any], vtt: bool = False) -> str:
    lines = ["WEBVTT", ""] if vtt else []
    for idx, segment in enumerate(project.get("segments", []), start=1):
        if not vtt:
            lines.append(str(idx))
        start = format_timestamp(segment.get("start", 0), "." if vtt else ",")
        end = format_timestamp(segment.get("end", 0), "." if vtt else ",")
        lines.append(f"{start} --> {end}")
        lines.append(f"{speaker_name(project, segment.get('speaker', 'SPEAKER_00'))}: {segment.get('text', '').strip()}")
        lines.append("")
    return "\n".join(lines)


def export_docx(project: dict[str, Any], output_path: Path, include_timestamps: bool = False) -> None:
    document = Document()
    document.add_heading(project.get("name") or "Transcripcion", level=1)
    if include_timestamps:
        for segment in project.get("segments", []):
            current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
            start = format_clock(segment.get("start", 0))
            end = format_clock(segment.get("end", 0))
            paragraph = document.add_paragraph()
            time_run = paragraph.add_run(f"[{start}-{end}] ")
            time_run.italic = True
            speaker_run = paragraph.add_run(f"{current}: ")
            speaker_run.bold = True
            paragraph.add_run((segment.get("text") or "").strip())
        document.save(output_path)
        return

    last_speaker = None
    paragraph = None
    for segment in project.get("segments", []):
        current = speaker_name(project, segment.get("speaker", "SPEAKER_00"))
        if current != last_speaker:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{current}: ")
            run.bold = True
            last_speaker = current
        if paragraph is None:
            paragraph = document.add_paragraph()
        paragraph.add_run((segment.get("text") or "").strip() + " ")
    document.save(output_path)
